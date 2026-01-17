"""
Test configuration and fixtures.
"""

import sys
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

# Add paths for imports
sys.path.insert(0, str(Path(__file__).parent.parent))
sys.path.insert(0, str(Path(__file__).parent.parent.parent / "src"))

from app.main import app


@pytest.fixture
def client():
    """Create test client."""
    return TestClient(app)


@pytest.fixture
def sample_docx_content():
    """Create a minimal DOCX file content for testing."""
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a test paragraph about SEO content optimization.")
    doc.add_heading("Section One", level=1)
    doc.add_paragraph(
        "Search engine optimization is important for visibility. "
        "Good content helps with rankings and user engagement."
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
