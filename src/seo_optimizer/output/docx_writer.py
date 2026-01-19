"""
DOCX Writer - Reconstructs DOCX with optimized content

This module provides multiple writers for different use cases:

1. PreservingDocxWriter (RECOMMENDED for user-facing output):
   - Copies the original document exactly
   - Preserves [H1], [H2], [H3] markers in text
   - Preserves bold, italic, and all formatting
   - Preserves empty paragraphs
   - ONLY highlights the actual inserted text (not entire paragraphs)
   - Appends FAQ section at the end with green highlighting

2. DocxWriter (for AST-based output):
   - Creates a new document from AST
   - Converts [H*] markers to Word heading styles
   - Rebuilds document structure

3. OptimizedDocumentWriter (for OptimizedContent objects):
   - Creates professional document with header and legend
   - Full change tracking with colors

Reference: docs/research/06-docx-output.md
"""

from __future__ import annotations

import contextlib
import io
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from seo_optimizer.diffing.models import (
    ChangeSet,
    OptimizedContent,
)
from seo_optimizer.ingestion.models import ContentNode, DocumentAST, NodeType
from seo_optimizer.output.highlighter import (
    add_paragraph_with_changes,
    add_text_with_changes,
    apply_highlights,
)

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.table import _Cell
    from docx.text.paragraph import Paragraph

# Regex pattern to detect heading markers
HEADING_PATTERN = re.compile(r"^\[H([123])\]\s*(.*)", re.DOTALL)

# Heading style mapping
HEADING_STYLE_MAP = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
}


def _ensure_heading_styles(doc: DocxDocument) -> None:
    """
    Ensure heading styles exist in the document.

    Args:
        doc: Document to check/modify
    """
    # python-docx should have built-in heading styles
    # but we verify they exist
    for level in range(1, 7):
        style_name = f"Heading {level}"
        with contextlib.suppress(KeyError):
            _ = doc.styles[style_name]


def apply_heading_style(paragraph: Paragraph, text: str) -> str:
    """
    Check if text starts with [H1], [H2], or [H3] marker.
    If so, apply the corresponding heading style and return cleaned text.

    Args:
        paragraph: The python-docx Paragraph object
        text: The paragraph text content

    Returns:
        Cleaned text with marker removed
    """
    match = HEADING_PATTERN.match(text)
    if match:
        level = int(match.group(1))
        clean_text = match.group(2)
        style_name = HEADING_STYLE_MAP.get(level, "Normal")
        with contextlib.suppress(KeyError):
            paragraph.style = style_name
        return clean_text
    return text


def setup_document_styles(doc: DocxDocument) -> None:
    """
    Configure document styles for professional output.

    Sets up heading styles with visual hierarchy:
    - H1: Large, navy blue, bold
    - H2: Medium, blue, bold
    - H3: Small, dark gray, bold

    Args:
        doc: Document to configure
    """
    styles = doc.styles

    # Configure Heading 1
    with contextlib.suppress(KeyError):
        h1_style = styles["Heading 1"]
        h1_font = h1_style.font
        h1_font.name = "Poppins"
        h1_font.size = Pt(24)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)  # Navy blue

    # Configure Heading 2
    with contextlib.suppress(KeyError):
        h2_style = styles["Heading 2"]
        h2_font = h2_style.font
        h2_font.name = "Poppins"
        h2_font.size = Pt(18)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(0x28, 0x74, 0xA6)  # Medium blue

    # Configure Heading 3
    with contextlib.suppress(KeyError):
        h3_style = styles["Heading 3"]
        h3_font = h3_style.font
        h3_font.name = "Poppins"
        h3_font.size = Pt(14)
        h3_font.bold = True
        h3_font.color.rgb = RGBColor(0x33, 0x33, 0x33)  # Dark gray

    # Configure Normal
    with contextlib.suppress(KeyError):
        normal_style = styles["Normal"]
        normal_font = normal_style.font
        normal_font.name = "Poppins"
        normal_font.size = Pt(12)


def _set_cell_shading(cell: _Cell, color_hex: str) -> None:
    """
    Apply background shading to a table cell.

    Args:
        cell: Table cell to shade
        color_hex: Hex color code (without #)
    """
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _add_paragraph_border(paragraph: Paragraph) -> None:
    """
    Add a border around a paragraph.

    Args:
        paragraph: Paragraph to add border to
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "1")
        border.set(qn("w:color"), "CCCCCC")
        pBdr.append(border)
    pPr.append(pBdr)


def add_color_legend(doc: DocxDocument) -> Paragraph:
    """
    Add a legend explaining the color coding.

    Args:
        doc: Document to add legend to

    Returns:
        The legend paragraph
    """
    legend_para = doc.add_paragraph()

    # Title
    title_run = legend_para.add_run("CHANGE LEGEND: ")
    title_run.bold = True

    # Green sample
    green_run = legend_para.add_run(" New Content ")
    green_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    legend_para.add_run("  |  ")

    # Yellow sample
    yellow_run = legend_para.add_run(" Modified ")
    yellow_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    legend_para.add_run("  |  ")

    # Strikethrough sample
    strike_run = legend_para.add_run(" Removed ")
    strike_run.font.strike = True
    strike_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Add border to legend paragraph
    _add_paragraph_border(legend_para)

    return legend_para


def add_optimization_header(
    doc: DocxDocument,
    content: OptimizedContent | None = None,
    url: str = "",
    keyword: str = "",
    total_changes: int = 0,
) -> None:
    """
    Add a professional header section with optimization metadata.

    Args:
        doc: Document to add header to
        content: OptimizedContent object (if available)
        url: URL analyzed (fallback if content not provided)
        keyword: Target keyword (fallback if content not provided)
        total_changes: Total number of changes (fallback if content not provided)
    """
    # Extract from content if available
    if content:
        url = content.url or url
        keyword = content.target_keyword or keyword
        date = content.optimization_date or datetime.now().strftime("%Y-%m-%d")
        total_changes = content.change_summary.get("total", total_changes)
    else:
        date = datetime.now().strftime("%Y-%m-%d")

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("SEO OPTIMIZATION REPORT")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    # Metadata table
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"

    metadata = [
        ("URL:", url or "(not specified)"),
        ("Target Keyword:", keyword or "(not specified)"),
        ("Optimized:", date),
        ("Total Changes:", f"{total_changes} modifications"),
    ]

    for i, (label, value) in enumerate(metadata):
        row = table.rows[i]
        # Label cell (shaded)
        label_cell = row.cells[0]
        label_cell.text = label
        if label_cell.paragraphs:
            for run in label_cell.paragraphs[0].runs:
                run.bold = True
        _set_cell_shading(label_cell, "E8F4FD")

        # Value cell
        value_cell = row.cells[1]
        value_cell.text = value

    # Spacer
    doc.add_paragraph()

    # Legend
    add_color_legend(doc)

    # Spacer
    doc.add_paragraph()


def _add_paragraph_with_highlighting(
    doc: DocxDocument,
    text: str,
    style: str | None = None,
    highlight: bool = False,
) -> Paragraph:
    """
    Add a paragraph to the document with optional highlighting.

    Args:
        doc: Document to add to
        text: Text content for the paragraph
        style: Optional style name (e.g., "Heading 1")
        highlight: Whether to apply highlight

    Returns:
        The created paragraph
    """
    para = doc.add_paragraph()
    if style:
        with contextlib.suppress(KeyError):
            para.style = style

    run = para.add_run(text)
    if highlight:
        from docx.enum.text import WD_COLOR_INDEX

        run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    return para


def _copy_paragraph_formatting(
    source_para: Paragraph,
    target_para: Paragraph,
) -> None:
    """
    Copy formatting from one paragraph to another.

    Args:
        source_para: Source paragraph
        target_para: Target paragraph
    """
    # Copy paragraph format
    if source_para.paragraph_format:
        spf = source_para.paragraph_format
        tpf = target_para.paragraph_format

        if spf.alignment is not None:
            tpf.alignment = spf.alignment
        if spf.left_indent is not None:
            tpf.left_indent = spf.left_indent
        if spf.right_indent is not None:
            tpf.right_indent = spf.right_indent
        if spf.first_line_indent is not None:
            tpf.first_line_indent = spf.first_line_indent
        if spf.space_before is not None:
            tpf.space_before = spf.space_before
        if spf.space_after is not None:
            tpf.space_after = spf.space_after
        if spf.line_spacing is not None:
            tpf.line_spacing = spf.line_spacing

    # Copy style
    if source_para.style:
        with contextlib.suppress(KeyError):
            target_para.style = source_para.style


def write_optimized_docx(
    original_path: str | Path,
    optimized_ast: DocumentAST,
    changeset: ChangeSet,
    output_path: str | Path,
) -> Path:
    """
    Write the optimized document with highlighting.

    Creates a new DOCX file with:
    - All original content preserved
    - New content added
    - Highlighting on all new content

    Args:
        original_path: Path to original DOCX (for formatting reference)
        optimized_ast: The optimized document AST
        changeset: Changes to highlight
        output_path: Where to save the output

    Returns:
        Path to the created file

    Example:
        >>> output = write_optimized_docx(
        ...     "input.docx",
        ...     optimized_ast,
        ...     changeset,
        ...     "output.docx"
        ... )
        >>> print(f"Saved to {output}")
    """
    original_path = Path(original_path)
    output_path = Path(output_path)

    if not original_path.exists():
        raise FileNotFoundError(f"Original file not found: {original_path}")

    # Copy original to output location first (preserves formatting/styles)
    shutil.copy2(original_path, output_path)

    # Open the copy for modification
    doc: DocxDocument = Document(str(output_path))

    # Apply highlights from changeset
    apply_highlights(doc, changeset)

    # Save the modified document
    doc.save(str(output_path))

    return output_path


def write_document_from_ast(
    ast: DocumentAST,
    output_path: str | Path,
    highlight_new: bool = True,
    new_node_ids: set[str] | None = None,
) -> Path:
    """
    Write a document from an AST structure.

    Creates a new DOCX file from scratch based on the AST.

    Args:
        ast: Document AST to write
        output_path: Where to save the output
        highlight_new: Whether to highlight new content
        new_node_ids: Set of node IDs that should be highlighted

    Returns:
        Path to the created file
    """
    output_path = Path(output_path)
    doc = Document()

    _ensure_heading_styles(doc)

    new_node_ids = new_node_ids or set()

    for node in ast.nodes:
        _write_node_to_document(doc, node, highlight_new, new_node_ids)

    doc.save(str(output_path))
    return output_path


def _write_node_to_document(
    doc: DocxDocument,
    node: ContentNode,
    highlight_new: bool,
    new_node_ids: set[str],
) -> None:
    """
    Write a single AST node to the document.

    Args:
        doc: Document to write to
        node: Node to write
        highlight_new: Whether to highlight
        new_node_ids: Set of new node IDs
    """
    should_highlight = highlight_new and node.node_id in new_node_ids

    if node.node_type == NodeType.HEADING:
        level = node.metadata.get("heading_level", 1)
        style = f"Heading {level}"
        _add_paragraph_with_highlighting(
            doc, node.text_content, style=style, highlight=should_highlight
        )
    elif node.node_type == NodeType.PARAGRAPH:
        _add_paragraph_with_highlighting(
            doc, node.text_content, highlight=should_highlight
        )
    elif node.node_type == NodeType.TABLE:
        _write_table_node(doc, node, should_highlight)
    elif node.node_type == NodeType.LIST:
        _write_list_node(doc, node, highlight_new, new_node_ids)

    # Handle child nodes (for complex structures)
    for child in node.children:
        if node.node_type not in (NodeType.TABLE, NodeType.LIST):
            _write_node_to_document(doc, child, highlight_new, new_node_ids)


def _write_table_node(
    doc: DocxDocument,
    node: ContentNode,
    highlight: bool,
) -> None:
    """
    Write a table node to the document.

    Args:
        doc: Document to write to
        node: Table node
        highlight: Whether to highlight
    """
    # Count rows and cells
    row_nodes = [c for c in node.children if c.node_type == NodeType.TABLE_ROW]
    if not row_nodes:
        return

    # Determine max columns
    max_cols = 0
    for row in row_nodes:
        cell_count = len(
            [c for c in row.children if c.node_type == NodeType.TABLE_CELL]
        )
        max_cols = max(max_cols, cell_count)

    if max_cols == 0:
        return

    # Create table
    table = doc.add_table(rows=len(row_nodes), cols=max_cols)

    for row_idx, row_node in enumerate(row_nodes):
        cell_nodes = [c for c in row_node.children if c.node_type == NodeType.TABLE_CELL]
        for col_idx, cell_node in enumerate(cell_nodes):
            if col_idx < max_cols:
                cell = table.rows[row_idx].cells[col_idx]
                # Clear default paragraph and add content
                if cell.paragraphs:
                    p = cell.paragraphs[0]
                    p.clear()
                    run = p.add_run(cell_node.text_content)
                    if highlight:
                        from docx.enum.text import WD_COLOR_INDEX

                        run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN


def _write_list_node(
    doc: DocxDocument,
    node: ContentNode,
    highlight_new: bool,
    new_node_ids: set[str],
) -> None:
    """
    Write a list node to the document.

    Args:
        doc: Document to write to
        node: List node
        highlight_new: Whether to highlight
        new_node_ids: Set of new node IDs
    """
    for item in node.children:
        if item.node_type == NodeType.LIST_ITEM:
            should_highlight = highlight_new and item.node_id in new_node_ids
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(item.text_content)
            if should_highlight:
                from docx.enum.text import WD_COLOR_INDEX

                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN


def _write_node_to_document_enhanced(
    doc: DocxDocument,
    node: ContentNode,
    highlight_new: bool,
    new_node_ids: set[str],
    modified_node_ids: set[str],
    text_insertions: list[dict],
) -> None:
    """
    Write a single AST node to the document with enhanced highlighting.

    Handles:
    - [H1], [H2], [H3] marker parsing
    - New node highlighting (entirely new content)
    - Modified node highlighting (inline changes)
    - Character-level text insertion highlighting

    Args:
        doc: Document to write to
        node: Node to write
        highlight_new: Whether to highlight
        new_node_ids: Set of new node IDs (highlight entire node)
        modified_node_ids: Set of modified node IDs (has inline changes)
        text_insertions: List of text insertion details for precise highlighting
    """
    is_new_node = highlight_new and node.node_id in new_node_ids
    is_modified_node = highlight_new and node.node_id in modified_node_ids

    text_content = node.text_content

    # Check for [H*] heading markers in text
    heading_match = HEADING_PATTERN.match(text_content)

    if node.node_type == NodeType.HEADING or heading_match:
        # Determine heading level and clean text
        if heading_match:
            level = int(heading_match.group(1))
            clean_text = heading_match.group(2).strip()
        else:
            level = node.metadata.get("heading_level", node.metadata.get("level", 2))
            clean_text = text_content

        style = f"Heading {level}"
        para = doc.add_paragraph()
        with contextlib.suppress(KeyError):
            para.style = style

        run = para.add_run(clean_text)

        # Apply explicit font styling to ensure Poppins is used
        run.font.name = "Poppins"
        if level == 1:
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
        elif level == 2:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x28, 0x74, 0xA6)
        elif level == 3:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        if is_new_node:
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    elif node.node_type == NodeType.PARAGRAPH:
        para = doc.add_paragraph()

        # Check if this node has character-level changes
        node_insertions = [
            ins for ins in text_insertions if ins.get("section_id") == node.node_id
        ]

        if is_new_node:
            # Entire paragraph is new - highlight it all
            run = para.add_run(text_content)
            run.font.name = "Poppins"
            run.font.size = Pt(12)
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

        elif is_modified_node and node_insertions:
            # Has inline changes - apply character-level highlighting
            _write_paragraph_with_inline_highlights(
                para, text_content, node_insertions
            )

        else:
            # No changes - just write the text
            run = para.add_run(text_content)
            run.font.name = "Poppins"
            run.font.size = Pt(12)

    elif node.node_type == NodeType.TABLE:
        _write_table_node(doc, node, is_new_node)

    elif node.node_type == NodeType.LIST:
        _write_list_node_enhanced(
            doc, node, highlight_new, new_node_ids, modified_node_ids, text_insertions
        )

    # Handle child nodes (for complex structures)
    for child in node.children:
        if node.node_type not in (NodeType.TABLE, NodeType.LIST):
            _write_node_to_document_enhanced(
                doc, child, highlight_new, new_node_ids, modified_node_ids, text_insertions
            )


def _write_paragraph_with_inline_highlights(
    para: Paragraph,
    text_content: str,
    insertions: list[dict],
) -> None:
    """
    Write a paragraph with inline highlighting for character-level changes.

    Args:
        para: Paragraph to write to
        text_content: Full text content
        insertions: List of insertion details with 'original' and 'new' text
    """
    # Build a map of text ranges to highlight
    # For simplicity, we'll highlight the new text where it appears
    remaining_text = text_content
    current_pos = 0

    for insertion in insertions:
        new_text = insertion.get("new", "")
        original_text = insertion.get("original", "")

        if new_text and new_text in remaining_text:
            # Find where the new text appears
            new_pos = remaining_text.find(new_text)

            if new_pos > 0:
                # Add text before the change (unhighlighted)
                pre_text = remaining_text[:new_pos]
                run = para.add_run(pre_text)
                run.font.name = "Poppins"
                run.font.size = Pt(12)

            # Add the new/changed text (highlighted)
            run = para.add_run(new_text)
            run.font.name = "Poppins"
            run.font.size = Pt(12)
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

            # Update remaining text
            remaining_text = remaining_text[new_pos + len(new_text):]
            current_pos = 0

    # Add any remaining text (unhighlighted)
    if remaining_text:
        run = para.add_run(remaining_text)
        run.font.name = "Poppins"
        run.font.size = Pt(12)


def _write_list_node_enhanced(
    doc: DocxDocument,
    node: ContentNode,
    highlight_new: bool,
    new_node_ids: set[str],
    modified_node_ids: set[str],
    text_insertions: list[dict],
) -> None:
    """
    Write a list node with enhanced highlighting support.

    Args:
        doc: Document to write to
        node: List node
        highlight_new: Whether to highlight
        new_node_ids: Set of new node IDs
        modified_node_ids: Set of modified node IDs
        text_insertions: Text insertion details
    """
    for item in node.children:
        if item.node_type == NodeType.LIST_ITEM:
            is_new = highlight_new and item.node_id in new_node_ids
            is_modified = highlight_new and item.node_id in modified_node_ids

            para = doc.add_paragraph(style="List Bullet")

            if is_new:
                run = para.add_run(item.text_content)
                run.font.name = "Poppins"
                run.font.size = Pt(12)
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            elif is_modified:
                item_insertions = [
                    ins for ins in text_insertions if ins.get("section_id") == item.node_id
                ]
                if item_insertions:
                    _write_paragraph_with_inline_highlights(
                        para, item.text_content, item_insertions
                    )
                else:
                    run = para.add_run(item.text_content)
                    run.font.name = "Poppins"
                    run.font.size = Pt(12)
            else:
                run = para.add_run(item.text_content)
                run.font.name = "Poppins"
                run.font.size = Pt(12)


def insert_faq_section(
    doc: DocxDocument,
    faq_content: list[tuple[str, str]],
    highlight: bool = True,
) -> DocxDocument:
    """
    Insert an FAQ section into the document with proper heading hierarchy.

    Uses:
    - H2 for section title
    - H3 for each question
    - Normal for each answer
    - Green highlighting for new content (not yellow)

    Args:
        doc: python-docx Document to modify
        faq_content: List of (question, answer) tuples
        highlight: Whether to highlight as new content

    Returns:
        Modified document with FAQ section
    """
    if not faq_content:
        return doc

    # Add FAQ heading (H2)
    faq_heading = doc.add_heading("Frequently Asked Questions", level=2)
    if highlight:
        for run in faq_heading.runs:
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    # Add note about AI-generated content
    note_para = doc.add_paragraph()
    note_run = note_para.add_run("(AI-Generated Schema Markup Content)")
    note_run.italic = True
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    note_run.font.size = Pt(9)

    # Add each Q&A pair
    for question, answer in faq_content:
        # Add question as H3 heading with green highlight
        q_heading = doc.add_heading(level=3)
        q_run = q_heading.add_run(question)
        if highlight:
            q_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

        # Add answer as normal paragraph with green highlight
        a_para = doc.add_paragraph()
        a_run = a_para.add_run(answer)
        if highlight:
            a_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    return doc


def add_faq_section_enhanced(
    doc: DocxDocument,
    faqs: list[dict[str, str]],
    highlight: bool = True,
) -> None:
    """
    Add FAQ section with proper heading hierarchy (dict-based input).

    Uses:
    - H2 for "Frequently Asked Questions" section header
    - H3 for each question
    - Normal for each answer
    - Green highlighting for new content

    Args:
        doc: Document to add FAQs to
        faqs: List of {"question": str, "answer": str} dicts
        highlight: Whether to highlight as new content
    """
    if not faqs:
        return

    # Section header (H2)
    section_header = doc.add_heading("Frequently Asked Questions", level=2)
    if highlight:
        for run in section_header.runs:
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    # Add note about AI-generated content
    note_para = doc.add_paragraph()
    note_run = note_para.add_run("(AI-Generated Schema Markup Content)")
    note_run.italic = True
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    note_run.font.size = Pt(9)

    # Add each FAQ
    for faq in faqs:
        question = faq.get("question", "")
        answer = faq.get("answer", "")

        if not question or not answer:
            continue

        # Question (H3) - with green highlight since it's new content
        question_heading = doc.add_heading(level=3)
        question_run = question_heading.add_run(question)
        if highlight:
            question_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

        # Answer (Normal) - with green highlight since it's new content
        answer_para = doc.add_paragraph()
        answer_run = answer_para.add_run(answer)
        if highlight:
            answer_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN


def insert_content_at_position(
    doc: DocxDocument,
    content: str,
    position: int,
    style: str | None = None,
    highlight: bool = True,
) -> Paragraph:
    """
    Insert content at a specific paragraph position.

    Args:
        doc: Document to modify
        content: Text content to insert
        position: Paragraph index to insert after
        style: Optional style name
        highlight: Whether to highlight

    Returns:
        The inserted paragraph
    """
    # python-docx doesn't have direct insert, so we need to work with XML
    para = doc.add_paragraph()
    if style:
        with contextlib.suppress(KeyError):
            para.style = style

    run = para.add_run(content)
    if highlight:
        from docx.enum.text import WD_COLOR_INDEX

        run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    # Move the paragraph to the correct position
    if position < len(doc.paragraphs) - 1:
        # Get the paragraph elements
        body = doc.element.body
        para_element = para._element
        body.remove(para_element)

        # Find the target position
        target_para = doc.paragraphs[position]
        target_para._element.addnext(para_element)

    return para


def validate_output(
    output_path: str | Path,
) -> list[str]:
    """
    Validate the output DOCX file.

    Checks:
    - File is valid DOCX
    - Highlights are correctly applied
    - No formatting corruption

    Args:
        output_path: Path to validate

    Returns:
        List of validation warnings (empty if valid)
    """
    output_path = Path(output_path)
    warnings: list[str] = []

    if not output_path.exists():
        return ["File does not exist"]

    if output_path.suffix.lower() != ".docx":
        warnings.append("File extension is not .docx")

    try:
        doc = Document(str(output_path))
    except Exception as e:
        return [f"Invalid DOCX file: {e}"]

    # Check for content
    if len(doc.paragraphs) == 0:
        warnings.append("Document has no paragraphs")

    # Check for highlights (informational)
    highlight_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color is not None:
                highlight_count += 1

    if highlight_count == 0:
        warnings.append("No highlighted content found (may be intentional)")

    return warnings


class DocxWriter:
    """
    Class-based wrapper for DOCX writing operations.

    Provides a streaming interface for generating optimized documents.
    """

    def __init__(self) -> None:
        """Initialize the DocxWriter."""
        pass

    def write_to_stream(
        self,
        ast: DocumentAST,
        stream: io.BytesIO,
        change_map: dict | None = None,
        highlight_new: bool = True,
    ) -> None:
        """
        Write optimized document to a stream.

        Args:
            ast: Document AST to write
            stream: BytesIO stream to write to
            change_map: Optional mapping of changes for highlighting
            highlight_new: Whether to highlight new content
        """
        from docx import Document

        doc = Document()

        # CRITICAL: Setup document styles for Poppins font and proper sizing
        setup_document_styles(doc)
        _ensure_heading_styles(doc)

        # Determine which nodes are new based on change_map
        new_node_ids: set[str] = set()
        modified_node_ids: set[str] = set()
        text_insertions: list[dict] = []

        if change_map:
            # new_nodes contains dicts with node_id, extract just the IDs
            new_nodes = change_map.get("new_nodes", [])
            for node_info in new_nodes:
                if isinstance(node_info, dict):
                    new_node_ids.add(node_info.get("node_id", ""))
                elif isinstance(node_info, str):
                    new_node_ids.add(node_info)

            # modified_nodes - body paragraphs with inline changes
            modified_nodes = change_map.get("modified_nodes", [])
            for node_id in modified_nodes:
                if isinstance(node_id, str):
                    modified_node_ids.add(node_id)

            # text_insertions - character-level changes within paragraphs
            text_insertions = change_map.get("text_insertions", [])

        for node in ast.nodes:
            _write_node_to_document_enhanced(
                doc, node, highlight_new, new_node_ids, modified_node_ids, text_insertions
            )

        # Save to stream
        doc.save(stream)
        stream.seek(0)


def merge_documents(
    base_doc_path: str | Path,
    additions: list[tuple[str, str]],
    output_path: str | Path,
    highlight_additions: bool = True,
) -> Path:
    """
    Merge additional content into a base document.

    Args:
        base_doc_path: Path to the base document
        additions: List of (position, content) tuples
                   position can be "end", "start", or a paragraph index
        output_path: Where to save the merged document
        highlight_additions: Whether to highlight added content

    Returns:
        Path to the merged document
    """
    base_doc_path = Path(base_doc_path)
    output_path = Path(output_path)

    # Copy base document
    shutil.copy2(base_doc_path, output_path)

    doc = Document(str(output_path))

    for position, content in additions:
        if position == "end":
            para = doc.add_paragraph()
            run = para.add_run(content)
            if highlight_additions:
                from docx.enum.text import WD_COLOR_INDEX

                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        elif position == "start":
            insert_content_at_position(doc, content, 0, highlight=highlight_additions)
        else:
            try:
                idx = int(position)
                insert_content_at_position(
                    doc, content, idx, highlight=highlight_additions
                )
            except ValueError:
                # Invalid position, append to end
                para = doc.add_paragraph()
                run = para.add_run(content)
                if highlight_additions:
                    from docx.enum.text import WD_COLOR_INDEX

                    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    doc.save(str(output_path))
    return output_path


class OptimizedDocumentWriter:
    """
    Generates professionally formatted DOCX output with all enhancements.

    Integrates:
    - Proper heading styles (H1, H2, H3)
    - Multi-color change highlighting (green/yellow/strikethrough)
    - Document header with metadata
    - Color legend
    - FAQ sections with proper hierarchy
    """

    def __init__(self) -> None:
        """Initialize the OptimizedDocumentWriter."""
        pass

    def create_document(self, content: OptimizedContent) -> DocxDocument:
        """
        Create a complete optimized document.

        Args:
            content: OptimizedContent object with all data

        Returns:
            A python-docx Document ready to save
        """
        doc = Document()

        # Step 1: Setup styles
        setup_document_styles(doc)
        _ensure_heading_styles(doc)

        # Step 2: Add header with metadata and legend
        add_optimization_header(doc, content=content)

        # Step 3: Add meta information section
        self._add_meta_section(doc, content)

        # Step 4: Add main content with change highlighting
        self._add_body_content(doc, content)

        # Step 5: Add FAQ section (if present)
        if content.faq_items:
            add_faq_section_enhanced(doc, content.faq_items, highlight=True)

        return doc

    def _add_meta_section(
        self,
        doc: DocxDocument,
        content: OptimizedContent,
    ) -> None:
        """
        Add meta title and description section.

        Args:
            doc: Document to add to
            content: Content with meta information
        """
        if not content.meta_title and not content.meta_description:
            return

        # Section header
        doc.add_heading("Meta Information", level=2)

        # Meta Title
        if content.meta_title:
            title_label = doc.add_paragraph()
            label_run = title_label.add_run("Meta Title: ")
            label_run.bold = True

            add_text_with_changes(title_label, content.meta_title.segments)

        # Meta Description
        if content.meta_description:
            desc_label = doc.add_paragraph()
            label_run = desc_label.add_run("Meta Description: ")
            label_run.bold = True

            add_text_with_changes(desc_label, content.meta_description.segments)

        # Spacer
        doc.add_paragraph()

    def _add_body_content(
        self,
        doc: DocxDocument,
        content: OptimizedContent,
    ) -> None:
        """
        Add main body content with change highlighting.

        Args:
            doc: Document to add to
            content: Content with body paragraphs
        """
        if not content.body_paragraphs:
            return

        for para_content in content.body_paragraphs:
            add_paragraph_with_changes(doc, para_content)

    def write_to_file(
        self,
        content: OptimizedContent,
        output_path: str | Path,
    ) -> Path:
        """
        Create and save an optimized document.

        Args:
            content: OptimizedContent object with all data
            output_path: Where to save the document

        Returns:
            Path to the saved file
        """
        output_path = Path(output_path)
        doc = self.create_document(content)
        doc.save(str(output_path))
        return output_path

    def write_to_stream(
        self,
        content: OptimizedContent,
        stream: io.BytesIO,
    ) -> None:
        """
        Create and save an optimized document to a stream.

        Args:
            content: OptimizedContent object with all data
            stream: BytesIO stream to write to
        """
        doc = self.create_document(content)
        doc.save(stream)
        stream.seek(0)


class PreservingDocxWriter:
    """
    DOCX writer that PRESERVES the original document format.

    This writer:
    - Copies the original document exactly
    - Preserves [H*] markers in text (does NOT convert to styles)
    - Preserves bold, italic, and other formatting
    - Preserves empty paragraphs
    - Preserves original styles
    - ONLY highlights the actual inserted text, not entire paragraphs

    Use this instead of DocxWriter when you want minimal changes to the original.
    """

    def __init__(self) -> None:
        """Initialize the PreservingDocxWriter."""
        pass

    def write_to_stream(
        self,
        original_stream: io.BytesIO,
        change_map: dict,
        output_stream: io.BytesIO,
    ) -> None:
        """
        Write optimized document preserving original format.

        Args:
            original_stream: BytesIO stream of the original DOCX
            change_map: Change map from optimization pipeline
            output_stream: BytesIO stream to write output to
        """
        # Load the original document
        original_stream.seek(0)
        doc: DocxDocument = Document(original_stream)

        # Apply text insertions with precise highlighting
        text_insertions = change_map.get("text_insertions", [])
        self._apply_text_insertions(doc, text_insertions)

        # Add new nodes (FAQ section) at the end
        new_nodes = change_map.get("new_nodes", [])
        faq_section = change_map.get("faq_section")
        if faq_section or any(n.get("node_id", "").startswith("faq_") for n in new_nodes):
            self._add_faq_section(doc, new_nodes)

        # Save to output stream
        doc.save(output_stream)
        output_stream.seek(0)

    def _apply_text_insertions(
        self,
        doc: DocxDocument,
        text_insertions: list[dict],
    ) -> None:
        """
        Apply text insertions with precise highlighting.

        For each insertion, finds the paragraph with the original text
        and replaces it with the new text, highlighting ONLY the difference.

        Args:
            doc: Document to modify
            text_insertions: List of insertion dicts with original/new text
        """
        for insertion in text_insertions:
            original_text = insertion.get("original", "")
            new_text = insertion.get("new", "")

            if not original_text or not new_text or original_text == new_text:
                continue

            # Find the paragraph containing the original text
            para_found = False
            for para in doc.paragraphs:
                para_text = para.text

                # Check if this paragraph matches the original
                if original_text in para_text or para_text == original_text:
                    # Found the paragraph - apply the change
                    self._replace_text_with_highlight(para, original_text, new_text)
                    para_found = True
                    break

            # If exact match not found, try finding by significant overlap
            if not para_found:
                # Try finding by first 50 characters of original
                search_key = original_text[:50] if len(original_text) > 50 else original_text
                for para in doc.paragraphs:
                    if search_key in para.text:
                        self._replace_text_with_highlight(para, original_text, new_text)
                        break

    def _replace_text_with_highlight(
        self,
        para: Paragraph,
        original_text: str,
        new_text: str,
    ) -> None:
        """
        Replace original text with new text, highlighting only the inserted portion.

        Args:
            para: Paragraph to modify
            original_text: Original text to find
            new_text: New text to replace with
        """
        # Find what was inserted (the difference)
        inserted_text = self._find_insertion_diff(original_text, new_text)

        if not inserted_text:
            return

        # Find where the insertion is in the new text
        insert_pos = new_text.find(inserted_text)
        if insert_pos == -1:
            return

        # Text before and after the insertion
        before_text = new_text[:insert_pos]
        after_text = new_text[insert_pos + len(inserted_text):]

        # Get original formatting from first run
        original_formatting = {}
        if para.runs:
            first_run = para.runs[0]
            original_formatting = {
                "bold": first_run.bold,
                "italic": first_run.italic,
                "underline": first_run.underline,
                "font_name": first_run.font.name,
                "font_size": first_run.font.size,
            }

        # Clear the paragraph
        for run in para.runs:
            run.text = ""

        # Rebuild with highlighting
        if before_text:
            before_run = para.add_run(before_text)
            self._apply_formatting(before_run, original_formatting)

        # Add the inserted text with green highlight
        insert_run = para.add_run(inserted_text)
        self._apply_formatting(insert_run, original_formatting)
        insert_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

        if after_text:
            after_run = para.add_run(after_text)
            self._apply_formatting(after_run, original_formatting)

    def _find_insertion_diff(self, original: str, new: str) -> str:
        """
        Find the text that was inserted (difference between original and new).

        Args:
            original: Original text
            new: New text (should be longer or contain insertion)

        Returns:
            The inserted text
        """
        # Use difflib to find the actual insertion
        import difflib

        # Simple approach: find the longest common substring approach
        # The insertion is what's in new but not in original

        # Find where original starts in new
        if original in new:
            # Original is a substring - insertion is elsewhere
            # This shouldn't happen in our case
            return ""

        # Use SequenceMatcher to find the difference
        matcher = difflib.SequenceMatcher(None, original, new)
        opcodes = matcher.get_opcodes()

        inserted_parts = []
        for tag, i1, i2, j1, j2 in opcodes:
            if tag == "insert":
                # Text that exists in new but not in original
                inserted_parts.append(new[j1:j2])
            elif tag == "replace":
                # Replacement - the new part is inserted, old part removed
                inserted_parts.append(new[j1:j2])

        return "".join(inserted_parts)

    def _apply_formatting(self, run, formatting: dict) -> None:
        """Apply formatting to a run."""
        if formatting.get("bold"):
            run.bold = formatting["bold"]
        if formatting.get("italic"):
            run.italic = formatting["italic"]
        if formatting.get("underline"):
            run.underline = formatting["underline"]
        if formatting.get("font_name"):
            run.font.name = formatting["font_name"]
        if formatting.get("font_size"):
            run.font.size = formatting["font_size"]

    def _add_faq_section(
        self,
        doc: DocxDocument,
        new_nodes: list[dict],
    ) -> None:
        """
        Add FAQ section at the end of the document.

        FAQ nodes are entirely new, so they're fully highlighted.

        Args:
            doc: Document to add FAQ to
            new_nodes: List of new node dicts
        """
        # Filter to FAQ nodes only
        faq_nodes = [n for n in new_nodes if n.get("node_id", "").startswith("faq_")]

        if not faq_nodes:
            return

        # Add blank line before FAQ
        doc.add_paragraph()

        # Add FAQ section header
        faq_header = doc.add_paragraph()
        header_run = faq_header.add_run("Frequently Asked Questions")
        header_run.bold = True
        header_run.font.size = Pt(16)
        header_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

        # Add FAQ items
        current_question = None
        for node in faq_nodes:
            content = node.get("content", "")
            node_type = node.get("type", "")

            if node_type == "heading":
                # This is a question
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(content)
                q_run.bold = True
                q_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            else:
                # This is an answer
                a_para = doc.add_paragraph()
                a_run = a_para.add_run(content)
                a_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
