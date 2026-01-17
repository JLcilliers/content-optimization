"""
Highlighter - Applies green highlighting to new content

Applies precise green highlighting to DOCX content using python-docx.
Handles run-level splitting for character-accurate highlighting.

Key implementation details:
- Uses WD_COLOR_INDEX.BRIGHT_GREEN (value 4) or YELLOW (value 7)
- Splits runs when highlighting partial content
- Preserves all original formatting

Reference: docs/research/06-docx-output.md
"""

from __future__ import annotations

from typing import TYPE_CHECKING, Any

from docx.enum.text import WD_COLOR_INDEX

from seo_optimizer.diffing.models import ChangeSet, HighlightRegion

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

# Type alias for python-docx Document (actual type at runtime)
# Using Any because docx.Document is a function, not a class type
DocxDocument = Any

# Highlight color index - YELLOW is more visible than bright green
# BRIGHT_GREEN = 4, YELLOW = 7
HIGHLIGHT_COLOR_INDEX = WD_COLOR_INDEX.YELLOW


def _copy_run_formatting(source_run: Run, target_run: Run) -> None:
    """
    Copy all formatting from one run to another.

    Args:
        source_run: Run to copy formatting from
        target_run: Run to copy formatting to
    """
    # Copy font properties
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.strike = source_run.font.strike
    target_run.font.subscript = source_run.font.subscript
    target_run.font.superscript = source_run.font.superscript

    # Copy font appearance
    if source_run.font.name:
        target_run.font.name = source_run.font.name
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb

    # Copy style if present
    if source_run.style:
        target_run.style = source_run.style


def _apply_highlight_to_run(
    run: Run,
    color: WD_COLOR_INDEX = HIGHLIGHT_COLOR_INDEX,
) -> None:
    """
    Apply highlight color to a run.

    Args:
        run: Run to highlight
        color: WD_COLOR_INDEX value (default: YELLOW)
    """
    run.font.highlight_color = color


def _get_paragraph_text_with_positions(
    paragraph: Paragraph,
) -> list[tuple[int, int, int, Run]]:
    """
    Get text positions for each run in a paragraph.

    Returns:
        List of (run_index, start_char, end_char, run) tuples
    """
    positions: list[tuple[int, int, int, Run]] = []
    current_pos = 0

    for idx, run in enumerate(paragraph.runs):
        run_text = run.text or ""
        run_len = len(run_text)
        if run_len > 0:
            positions.append((idx, current_pos, current_pos + run_len, run))
        current_pos += run_len

    return positions


def split_run_for_highlight(
    paragraph: Paragraph,
    run_index: int,
    start_char: int,
    end_char: int,
) -> list[Run]:
    """
    Split a run to enable partial highlighting.

    When a highlight region falls within a run (not at boundaries),
    the run must be split into multiple runs:
    - Pre-highlight portion (no highlight)
    - Highlighted portion (with highlight)
    - Post-highlight portion (no highlight)

    Args:
        paragraph: The paragraph containing the run
        run_index: Index of the run to split
        start_char: Start position within the run (relative to run start)
        end_char: End position within the run (relative to run start)

    Returns:
        List of new runs created (the highlighted run is in the middle)

    Example:
        Original run: "Hello world"
        Highlight "world" (positions 6-11)
        Result: ["Hello ", "world"] with second run highlighted
    """
    run = paragraph.runs[run_index]
    original_text = run.text or ""

    if not original_text:
        return []

    # Validate positions
    start_char = max(0, start_char)
    end_char = min(len(original_text), end_char)

    if start_char >= end_char:
        return []

    # Split text into parts
    pre_text = original_text[:start_char]
    highlight_text = original_text[start_char:end_char]
    post_text = original_text[end_char:]

    new_runs: list[Run] = []

    # Get the run's XML element for insertion
    run_element = run._element

    # Modify the original run to contain only pre-text (or highlight if no pre)
    if pre_text:
        run.text = pre_text
        new_runs.append(run)

        # Create new run for highlighted portion
        highlight_run = paragraph.add_run(highlight_text)
        _copy_run_formatting(run, highlight_run)
        _apply_highlight_to_run(highlight_run)

        # Move the new run to the correct position (after original)
        run_element.addnext(highlight_run._element)
        new_runs.append(highlight_run)

        # Create run for post-highlight if needed
        if post_text:
            post_run = paragraph.add_run(post_text)
            _copy_run_formatting(run, post_run)
            highlight_run._element.addnext(post_run._element)
            new_runs.append(post_run)
    else:
        # No pre-text - modify original to be highlighted
        run.text = highlight_text
        _apply_highlight_to_run(run)
        new_runs.append(run)

        # Create run for post-highlight if needed
        if post_text:
            post_run = paragraph.add_run(post_text)
            _copy_run_formatting(run, post_run)
            run_element.addnext(post_run._element)
            new_runs.append(post_run)

    return new_runs


def _find_paragraph_by_position(
    document: DocxDocument,
    node_id: str,
    position_map: dict[str, int] | None = None,
) -> Paragraph | None:
    """
    Find a paragraph in the document by node ID or position.

    Args:
        document: The document to search
        node_id: Node ID to find (format: doc_id_p0, doc_id_h1, etc.)
        position_map: Optional mapping of node_id to paragraph index

    Returns:
        The paragraph or None if not found
    """
    # Extract position info from node_id
    # Format: doc_xxxxxxxx_p0 or doc_xxxxxxxx_h1
    parts = node_id.split("_")
    if len(parts) < 3:
        return None

    pos_part = parts[-1]  # e.g., "p0", "h1"

    if position_map and node_id in position_map:
        idx = position_map[node_id]
        if 0 <= idx < len(document.paragraphs):
            para: Paragraph = document.paragraphs[idx]
            return para

    # Try to extract index from position part
    if pos_part.startswith("p") or pos_part.startswith("h"):
        try:
            idx = int(pos_part[1:])
            if 0 <= idx < len(document.paragraphs):
                para2: Paragraph = document.paragraphs[idx]
                return para2
        except ValueError:
            pass

    return None


def highlight_region(
    document: DocxDocument,
    region: HighlightRegion,
    paragraph_index: int | None = None,
) -> bool:
    """
    Highlight a single region in the document.

    Handles run splitting if the region doesn't align
    with existing run boundaries.

    Args:
        document: Document to modify
        region: Region to highlight
        paragraph_index: Optional direct paragraph index

    Returns:
        True if highlighting was successful, False otherwise
    """
    # Find the paragraph
    paragraph = None
    if paragraph_index is not None:
        if 0 <= paragraph_index < len(document.paragraphs):
            paragraph = document.paragraphs[paragraph_index]
    else:
        paragraph = _find_paragraph_by_position(document, region.node_id)

    if paragraph is None:
        return False

    # Get run positions in paragraph
    run_positions = _get_paragraph_text_with_positions(paragraph)
    if not run_positions:
        return False

    # Find which runs the region spans
    region_start = region.start_char
    region_end = region.end_char

    # Process runs that overlap with the region
    runs_to_split: list[tuple[int, int, int]] = []  # (run_idx, local_start, local_end)

    for run_idx, run_start, run_end, _run in run_positions:
        # Check if this run overlaps with the region
        if run_end <= region_start or run_start >= region_end:
            continue  # No overlap

        # Calculate local positions within this run
        local_start = max(0, region_start - run_start)
        local_end = min(run_end - run_start, region_end - run_start)

        runs_to_split.append((run_idx, local_start, local_end))

    # Apply highlighting (process in reverse to maintain indices)
    for run_idx, local_start, local_end in reversed(runs_to_split):
        run = paragraph.runs[run_idx]
        run_text = run.text or ""

        # Check if we need to split or can highlight the whole run
        if local_start == 0 and local_end == len(run_text):
            # Highlight entire run
            _apply_highlight_to_run(run)
        else:
            # Need to split
            split_run_for_highlight(paragraph, run_idx, local_start, local_end)

    return True


def apply_highlights(
    document: DocxDocument,
    changeset: ChangeSet,
) -> DocxDocument:
    """
    Apply highlighting to all additions in the changeset.

    Modifies the document in place, applying yellow highlighting
    to all regions identified in the changeset.

    Args:
        document: python-docx Document object
        changeset: ChangeSet with regions to highlight

    Returns:
        Modified document with highlights applied

    CRITICAL:
        Only regions in the changeset should be highlighted.
        Original content must remain unchanged.

    Example:
        >>> from docx import Document
        >>> doc = Document("optimized.docx")
        >>> doc = apply_highlights(doc, changeset)
        >>> doc.save("output.docx")
    """
    if not changeset.additions:
        return document

    # Process each addition
    for addition in changeset.additions:
        # Only process high-confidence additions
        if addition.confidence < 0.7:
            continue

        for region in addition.highlight_regions:
            # Try to find and highlight the region
            highlight_region(document, region)

    return document


def highlight_text_in_paragraph(
    paragraph: Paragraph,
    text_to_highlight: str,
    case_sensitive: bool = False,
) -> bool:
    """
    Find and highlight specific text within a paragraph.

    This is a convenience function for highlighting by text content
    rather than character positions.

    Args:
        paragraph: Paragraph to search and modify
        text_to_highlight: Text to find and highlight
        case_sensitive: Whether search is case-sensitive

    Returns:
        True if text was found and highlighted, False otherwise
    """
    full_text = "".join(run.text or "" for run in paragraph.runs)

    search_text = text_to_highlight if case_sensitive else text_to_highlight.lower()
    search_in = full_text if case_sensitive else full_text.lower()

    start_pos = search_in.find(search_text)
    if start_pos == -1:
        return False

    end_pos = start_pos + len(text_to_highlight)

    # Use a direct approach since we have the paragraph
    run_positions = _get_paragraph_text_with_positions(paragraph)

    for run_idx, run_start, run_end, run in run_positions:
        if run_end <= start_pos or run_start >= end_pos:
            continue

        local_start = max(0, start_pos - run_start)
        local_end = min(run_end - run_start, end_pos - run_start)

        run_text = run.text or ""
        if local_start == 0 and local_end == len(run_text):
            _apply_highlight_to_run(run)
        else:
            split_run_for_highlight(paragraph, run_idx, local_start, local_end)

    return True


def highlight_new_paragraph(paragraph: Paragraph) -> None:
    """
    Highlight an entire paragraph as new content.

    Use this for newly added paragraphs where all content should be highlighted.

    Args:
        paragraph: Paragraph to highlight entirely
    """
    for run in paragraph.runs:
        _apply_highlight_to_run(run)


def create_highlighted_run(
    paragraph: Paragraph,
    text: str,
    copy_formatting_from: Run | None = None,
) -> Run:
    """
    Create a new run with highlighting applied.

    Args:
        paragraph: Paragraph to add the run to
        text: Text content for the run
        copy_formatting_from: Optional run to copy formatting from

    Returns:
        The new highlighted run
    """
    new_run = paragraph.add_run(text)

    if copy_formatting_from:
        _copy_run_formatting(copy_formatting_from, new_run)

    _apply_highlight_to_run(new_run)
    return new_run
