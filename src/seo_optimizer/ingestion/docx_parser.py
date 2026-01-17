"""
DOCX Parser - Parses DOCX files into Document AST

This module handles parsing DOCX files using python-docx while:
- Preserving exact document structure (headings, lists, tables)
- Tracking character positions for diffing
- Extracting run-level formatting for reconstruction
- Creating OriginalSnapshot for change detection

Reference: docs/research/01-docx-parsing.md
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from seo_optimizer.ingestion.models import (
    ContentNode,
    DocumentAST,
    DocumentMetadata,
    FormattingInfo,
    NodeType,
    OriginalSnapshot,
    PositionInfo,
    TextRun,
)

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run


def _extract_run_formatting(run: Run) -> FormattingInfo:
    """
    Extract formatting information from a python-docx Run.

    Args:
        run: A python-docx Run object

    Returns:
        FormattingInfo with all formatting details
    """
    font = run.font

    # Get font size in points
    font_size = None
    if font.size is not None:
        font_size = font.size.pt

    # Get font color as hex
    font_color = None
    if font.color and font.color.rgb:
        font_color = str(font.color.rgb)

    return FormattingInfo(
        bold=bool(font.bold),
        italic=bool(font.italic),
        underline=bool(font.underline),
        strike=bool(font.strike),
        font_name=font.name,
        font_size=font_size,
        font_color=font_color,
        style_name=run.style.name if run.style else None,
    )


def _extract_paragraph_formatting(paragraph: Paragraph) -> FormattingInfo:
    """
    Extract paragraph-level formatting.

    Args:
        paragraph: A python-docx Paragraph object

    Returns:
        FormattingInfo with paragraph formatting
    """
    pf = paragraph.paragraph_format

    # Get alignment as string
    alignment = None
    if paragraph.alignment is not None:
        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
            WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify",
        }
        alignment = alignment_map.get(paragraph.alignment, "left")

    # Get spacing
    space_before = None
    space_after = None
    if pf.space_before is not None:
        space_before = pf.space_before.pt
    if pf.space_after is not None:
        space_after = pf.space_after.pt

    # Get line spacing
    line_spacing = None
    if pf.line_spacing is not None:
        if isinstance(pf.line_spacing, Pt):
            line_spacing = pf.line_spacing.pt
        else:
            line_spacing = float(pf.line_spacing)

    # Get indentation
    indent_left = None
    indent_right = None
    indent_first_line = None
    if pf.left_indent is not None:
        indent_left = pf.left_indent.inches
    if pf.right_indent is not None:
        indent_right = pf.right_indent.inches
    if pf.first_line_indent is not None:
        indent_first_line = pf.first_line_indent.inches

    # Get heading level from style
    heading_level = None
    if paragraph.style and paragraph.style.name:
        style_name = paragraph.style.name
        if style_name.startswith("Heading "):
            level_str = style_name.replace("Heading ", "")
            if level_str.isdigit():
                heading_level = int(level_str)

    return FormattingInfo(
        alignment=alignment,
        space_before=space_before,
        space_after=space_after,
        line_spacing=line_spacing,
        indent_left=indent_left,
        indent_right=indent_right,
        indent_first_line=indent_first_line,
        style_name=paragraph.style.name if paragraph.style else None,
        heading_level=heading_level,
    )


def _detect_heading_level(paragraph: Paragraph) -> int:
    """
    Detect if a paragraph is a heading and return its level.

    Args:
        paragraph: A python-docx Paragraph object

    Returns:
        Heading level (1-6) or 0 if not a heading
    """
    if paragraph.style and paragraph.style.name:
        style_name = paragraph.style.name
        if style_name.startswith("Heading "):
            try:
                return int(style_name.replace("Heading ", ""))
            except ValueError:
                pass
        # Also check for Title style
        if style_name == "Title":
            return 1
        if style_name == "Subtitle":
            return 2
    return 0


def _parse_paragraph(
    paragraph: Paragraph,
    para_index: int,
    char_offset: int,
    doc_id: str,
) -> tuple[ContentNode, int]:
    """
    Parse a paragraph into a ContentNode.

    Args:
        paragraph: python-docx Paragraph
        para_index: Index of this paragraph
        char_offset: Current character offset in document
        doc_id: Document ID for unique IDs

    Returns:
        Tuple of (ContentNode, new_char_offset)
    """
    heading_level = _detect_heading_level(paragraph)
    node_type = NodeType.HEADING if heading_level > 0 else NodeType.PARAGRAPH

    # Build position ID
    position_id = f"h{para_index}" if heading_level > 0 else f"p{para_index}"
    node_id = f"{doc_id}_{position_id}"

    # Extract text content and runs
    text_parts: list[str] = []
    runs: list[TextRun] = []
    run_offset = char_offset

    for run_idx, run in enumerate(paragraph.runs):
        run_text = run.text
        if not run_text:
            continue

        text_parts.append(run_text)

        # Create TextRun with formatting
        run_formatting = _extract_run_formatting(run)
        run_position = PositionInfo(
            position_id=f"{position_id}_r{run_idx}",
            start_char=run_offset,
            end_char=run_offset + len(run_text),
            parent_id=position_id,
        )

        # Check for hyperlink
        hyperlink_url = None
        # python-docx doesn't directly expose hyperlinks on runs
        # We'd need to access the underlying XML for full hyperlink support

        runs.append(
            TextRun(
                text=run_text,
                formatting=run_formatting,
                position=run_position,
                hyperlink_url=hyperlink_url,
            )
        )
        run_offset += len(run_text)

    text_content = "".join(text_parts)
    end_char = char_offset + len(text_content)

    # Create position info
    position = PositionInfo(
        position_id=position_id,
        start_char=char_offset,
        end_char=end_char,
    )

    # Build metadata
    metadata: dict[str, Any] = {}
    if heading_level > 0:
        metadata["heading_level"] = heading_level

    para_formatting = _extract_paragraph_formatting(paragraph)
    if para_formatting.style_name:
        metadata["style_name"] = para_formatting.style_name

    return (
        ContentNode(
            node_id=node_id,
            node_type=node_type,
            position=position,
            text_content=text_content,
            runs=runs,
            metadata=metadata,
        ),
        end_char,
    )


def _parse_table(
    table: Table,
    table_index: int,
    char_offset: int,
    doc_id: str,
) -> tuple[ContentNode, int]:
    """
    Parse a table into a ContentNode with nested row/cell nodes.

    Args:
        table: python-docx Table
        table_index: Index of this table
        char_offset: Current character offset
        doc_id: Document ID

    Returns:
        Tuple of (ContentNode, new_char_offset)
    """
    position_id = f"t{table_index}"
    node_id = f"{doc_id}_{position_id}"

    table_text_parts: list[str] = []
    row_nodes: list[ContentNode] = []
    current_offset = char_offset

    for row_idx, row in enumerate(table.rows):
        row_position_id = f"{position_id}_r{row_idx}"
        row_node_id = f"{doc_id}_{row_position_id}"

        cell_nodes: list[ContentNode] = []
        row_text_parts: list[str] = []
        row_start_char = current_offset

        for cell_idx, cell in enumerate(row.cells):
            cell_position_id = f"{row_position_id}_c{cell_idx}"
            cell_node_id = f"{doc_id}_{cell_position_id}"

            cell_text = cell.text
            cell_start_char = current_offset

            cell_node = ContentNode(
                node_id=cell_node_id,
                node_type=NodeType.TABLE_CELL,
                position=PositionInfo(
                    position_id=cell_position_id,
                    start_char=cell_start_char,
                    end_char=cell_start_char + len(cell_text),
                    parent_id=row_position_id,
                ),
                text_content=cell_text,
            )
            cell_nodes.append(cell_node)
            row_text_parts.append(cell_text)
            current_offset += len(cell_text) + 1  # +1 for cell separator

        row_text = "\t".join(row_text_parts)
        table_text_parts.append(row_text)

        row_node = ContentNode(
            node_id=row_node_id,
            node_type=NodeType.TABLE_ROW,
            position=PositionInfo(
                position_id=row_position_id,
                start_char=row_start_char,
                end_char=current_offset - 1,
                parent_id=position_id,
            ),
            text_content=row_text,
            children=cell_nodes,
        )
        row_nodes.append(row_node)
        current_offset += 1  # +1 for row separator

    table_text = "\n".join(table_text_parts)

    return (
        ContentNode(
            node_id=node_id,
            node_type=NodeType.TABLE,
            position=PositionInfo(
                position_id=position_id,
                start_char=char_offset,
                end_char=current_offset - 1,
            ),
            text_content=table_text,
            children=row_nodes,
        ),
        current_offset,
    )


def parse_docx(file_path: str | Path) -> DocumentAST:
    """
    Parse a DOCX file into a Document AST.

    Args:
        file_path: Path to the DOCX file to parse

    Returns:
        DocumentAST with full structure and formatting preserved

    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file is not a valid DOCX

    Example:
        >>> ast = parse_docx("document.docx")
        >>> print(f"Parsed {len(ast.nodes)} top-level elements")
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    if not path.suffix.lower() == ".docx":
        raise ValueError(f"Not a DOCX file: {file_path}")

    try:
        doc: DocxDocument = Document(str(path))
    except Exception as e:
        raise ValueError(f"Invalid DOCX file: {file_path}. Error: {e}") from e

    # Generate document ID
    doc_id = f"doc_{uuid4().hex[:8]}"

    # Extract document metadata
    core_props = doc.core_properties
    metadata = DocumentMetadata(
        source_path=str(path.absolute()),
        file_size=path.stat().st_size,
        title=core_props.title,
        author=core_props.author,
        created=core_props.created,
        modified=core_props.modified,
    )

    # Parse document body
    nodes: list[ContentNode] = []
    full_text_parts: list[str] = []
    char_offset = 0
    para_index = 0
    table_index = 0

    # Iterate through document body elements
    for element in doc.element.body:
        # Check if it's a paragraph
        if element.tag.endswith("}p"):
            # Find the corresponding paragraph object
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break

            if para is not None:
                node, char_offset = _parse_paragraph(
                    para, para_index, char_offset, doc_id
                )
                if node.text_content:  # Skip empty paragraphs
                    nodes.append(node)
                    full_text_parts.append(node.text_content)
                    char_offset += 1  # Add separator
                para_index += 1

        # Check if it's a table
        elif element.tag.endswith("}tbl"):
            # Find the corresponding table object
            tbl = None
            for t in doc.tables:
                if t._element is element:
                    tbl = t
                    break

            if tbl is not None:
                node, char_offset = _parse_table(tbl, table_index, char_offset, doc_id)
                nodes.append(node)
                full_text_parts.append(node.text_content)
                char_offset += 1  # Add separator
                table_index += 1

    full_text = "\n".join(full_text_parts)

    return DocumentAST(
        doc_id=doc_id,
        nodes=nodes,
        metadata=metadata,
        full_text=full_text,
        char_count=len(full_text),
    )


def create_snapshot(ast: DocumentAST) -> OriginalSnapshot:
    """
    Create an immutable snapshot from a Document AST.

    The snapshot is used as the baseline for diffing to detect
    what content is genuinely new vs. existing.

    Args:
        ast: The Document AST to snapshot

    Returns:
        Frozen OriginalSnapshot for diffing baseline

    Example:
        >>> ast = parse_docx("document.docx")
        >>> snapshot = create_snapshot(ast)
        >>> # Snapshot is now immutable
    """
    return OriginalSnapshot.from_document_ast(ast)


def parse_docx_with_snapshot(file_path: str | Path) -> tuple[DocumentAST, OriginalSnapshot]:
    """
    Parse a DOCX file and immediately create a snapshot.

    This is the recommended entry point for the optimization pipeline
    as it ensures the snapshot is created before any modifications.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Tuple of (DocumentAST, OriginalSnapshot)

    Example:
        >>> ast, snapshot = parse_docx_with_snapshot("input.docx")
        >>> # snapshot captures original state for diffing
    """
    ast = parse_docx(file_path)
    snapshot = create_snapshot(ast)
    return ast, snapshot
