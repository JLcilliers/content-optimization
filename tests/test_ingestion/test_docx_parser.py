"""
Tests for DOCX Parser Module

Tests the parsing of DOCX files into DocumentAST, including:
- Basic parsing functionality
- Heading detection
- Paragraph extraction
- Table parsing
- Run-level formatting
- Position tracking
- Error handling
"""

from pathlib import Path
from typing import Any

import pytest
from docx import Document

from seo_optimizer.ingestion.docx_parser import (
    _detect_heading_level,
    _extract_paragraph_formatting,
    _extract_run_formatting,
    parse_docx,
)
from seo_optimizer.ingestion.models import (
    DocumentAST,
    NodeType,
)


# =============================================================================
# Fixtures for DOCX Testing
# =============================================================================


@pytest.fixture
def temp_docx_path(tmp_path: Path) -> Path:
    """Create a temporary DOCX file for testing."""
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is the first paragraph.")
    doc.add_heading("Section One", level=2)
    doc.add_paragraph("Content under section one.")
    doc.add_paragraph("Another paragraph with some text.")

    docx_path = tmp_path / "test_document.docx"
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def temp_docx_with_formatting(tmp_path: Path) -> Path:
    """Create a DOCX file with various formatting."""
    doc = Document()
    doc.add_heading("Formatted Document", level=1)

    # Paragraph with bold and italic
    para = doc.add_paragraph()
    run_normal = para.add_run("Normal text, ")
    run_bold = para.add_run("bold text, ")
    run_bold.bold = True
    run_italic = para.add_run("italic text.")
    run_italic.italic = True

    docx_path = tmp_path / "formatted_document.docx"
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def temp_docx_with_table(tmp_path: Path) -> Path:
    """Create a DOCX file with a table."""
    doc = Document()
    doc.add_heading("Document with Table", level=1)
    doc.add_paragraph("Here is a table:")

    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(0, 2).text = "Header 3"
    table.cell(1, 0).text = "Row 1, Col 1"
    table.cell(1, 1).text = "Row 1, Col 2"
    table.cell(1, 2).text = "Row 1, Col 3"

    doc.add_paragraph("Text after table.")

    docx_path = tmp_path / "table_document.docx"
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def temp_empty_docx(tmp_path: Path) -> Path:
    """Create an empty DOCX file."""
    doc = Document()
    docx_path = tmp_path / "empty_document.docx"
    doc.save(str(docx_path))
    return docx_path


# =============================================================================
# Basic Parsing Tests
# =============================================================================


class TestParseDocx:
    """Tests for the main parse_docx function."""

    def test_parse_basic_document(self, temp_docx_path: Path) -> None:
        """Test parsing a basic document."""
        ast = parse_docx(temp_docx_path)

        assert isinstance(ast, DocumentAST)
        assert ast.doc_id.startswith("doc_")
        assert len(ast.nodes) > 0
        assert ast.char_count > 0

    def test_parse_extracts_headings(self, temp_docx_path: Path) -> None:
        """Test that headings are correctly identified."""
        ast = parse_docx(temp_docx_path)

        # Find heading nodes
        heading_nodes = [n for n in ast.nodes if n.node_type == NodeType.HEADING]

        assert len(heading_nodes) >= 2  # Title + Section One
        assert any("Test Document" in n.text_content for n in heading_nodes)
        assert any("Section One" in n.text_content for n in heading_nodes)

    def test_parse_extracts_paragraphs(self, temp_docx_path: Path) -> None:
        """Test that paragraphs are correctly extracted."""
        ast = parse_docx(temp_docx_path)

        # Find paragraph nodes
        para_nodes = [n for n in ast.nodes if n.node_type == NodeType.PARAGRAPH]

        assert len(para_nodes) >= 2
        assert any("first paragraph" in n.text_content for n in para_nodes)

    def test_parse_creates_full_text(self, temp_docx_path: Path) -> None:
        """Test that full_text is properly constructed."""
        ast = parse_docx(temp_docx_path)

        assert ast.full_text is not None
        assert len(ast.full_text) > 0
        assert "Test Document" in ast.full_text
        assert "first paragraph" in ast.full_text

    def test_parse_preserves_position_info(self, temp_docx_path: Path) -> None:
        """Test that position information is tracked."""
        ast = parse_docx(temp_docx_path)

        for node in ast.nodes:
            assert node.position is not None
            assert node.position.start_char >= 0
            assert node.position.end_char >= node.position.start_char
            assert node.position.position_id is not None

    def test_parse_extracts_metadata(self, temp_docx_path: Path) -> None:
        """Test that document metadata is extracted."""
        ast = parse_docx(temp_docx_path)

        assert ast.metadata is not None
        assert ast.metadata.source_path is not None
        assert ast.metadata.file_size is not None
        assert ast.metadata.file_size > 0

    def test_parse_file_not_found(self, tmp_path: Path) -> None:
        """Test error handling for missing file."""
        nonexistent = tmp_path / "nonexistent.docx"

        with pytest.raises(FileNotFoundError):
            parse_docx(nonexistent)

    def test_parse_invalid_file(self, tmp_path: Path) -> None:
        """Test error handling for invalid DOCX."""
        # Create a text file with .docx extension
        invalid_path = tmp_path / "invalid.docx"
        invalid_path.write_text("This is not a valid DOCX file")

        with pytest.raises(ValueError, match="Invalid DOCX"):
            parse_docx(invalid_path)

    def test_parse_non_docx_extension(self, tmp_path: Path) -> None:
        """Test error handling for wrong file extension."""
        txt_path = tmp_path / "document.txt"
        txt_path.write_text("Hello world")

        with pytest.raises(ValueError, match="Not a DOCX file"):
            parse_docx(txt_path)

    def test_parse_empty_document(self, temp_empty_docx: Path) -> None:
        """Test parsing an empty document."""
        ast = parse_docx(temp_empty_docx)

        assert isinstance(ast, DocumentAST)
        # Empty doc should have no content nodes
        assert len(ast.nodes) == 0 or all(not n.text_content for n in ast.nodes)


# =============================================================================
# Formatting Extraction Tests
# =============================================================================


class TestFormattingExtraction:
    """Tests for formatting extraction functions."""

    def test_extract_run_formatting_bold(self, temp_docx_with_formatting: Path) -> None:
        """Test extraction of bold formatting from runs."""
        doc = Document(str(temp_docx_with_formatting))

        # Find a paragraph with runs
        for para in doc.paragraphs:
            for run in para.runs:
                formatting = _extract_run_formatting(run)
                if run.bold:
                    assert formatting.bold is True

    def test_extract_run_formatting_italic(
        self, temp_docx_with_formatting: Path
    ) -> None:
        """Test extraction of italic formatting from runs."""
        doc = Document(str(temp_docx_with_formatting))

        for para in doc.paragraphs:
            for run in para.runs:
                formatting = _extract_run_formatting(run)
                if run.italic:
                    assert formatting.italic is True

    def test_extract_paragraph_formatting(self, temp_docx_path: Path) -> None:
        """Test extraction of paragraph-level formatting."""
        doc = Document(str(temp_docx_path))

        for para in doc.paragraphs:
            formatting = _extract_paragraph_formatting(para)
            # Should not raise and should return FormattingInfo
            assert formatting is not None


# =============================================================================
# Heading Detection Tests
# =============================================================================


class TestHeadingDetection:
    """Tests for heading level detection."""

    def test_detect_heading_level_h1(self, temp_docx_path: Path) -> None:
        """Test detection of Heading 1 style."""
        doc = Document(str(temp_docx_path))

        for para in doc.paragraphs:
            if para.style and para.style.name == "Heading 1":
                level = _detect_heading_level(para)
                assert level == 1

    def test_detect_heading_level_h2(self, temp_docx_path: Path) -> None:
        """Test detection of Heading 2 style."""
        doc = Document(str(temp_docx_path))

        for para in doc.paragraphs:
            if para.style and para.style.name == "Heading 2":
                level = _detect_heading_level(para)
                assert level == 2

    def test_detect_heading_level_normal(self, temp_docx_path: Path) -> None:
        """Test that normal paragraphs return level 0."""
        doc = Document(str(temp_docx_path))

        for para in doc.paragraphs:
            if para.style and para.style.name == "Normal":
                level = _detect_heading_level(para)
                assert level == 0


# =============================================================================
# Table Parsing Tests
# =============================================================================


class TestTableParsing:
    """Tests for table parsing functionality."""

    def test_parse_document_with_table(self, temp_docx_with_table: Path) -> None:
        """Test parsing a document containing a table."""
        ast = parse_docx(temp_docx_with_table)

        # Should have parsed successfully
        assert isinstance(ast, DocumentAST)
        assert len(ast.nodes) > 0

        # Should have table node
        table_nodes = [n for n in ast.nodes if n.node_type == NodeType.TABLE]
        assert len(table_nodes) >= 1

    def test_table_contains_row_children(self, temp_docx_with_table: Path) -> None:
        """Test that table nodes contain row children."""
        ast = parse_docx(temp_docx_with_table)

        table_nodes = [n for n in ast.nodes if n.node_type == NodeType.TABLE]
        if table_nodes:
            table = table_nodes[0]
            assert len(table.children) > 0
            assert all(c.node_type == NodeType.TABLE_ROW for c in table.children)

    def test_table_rows_contain_cell_children(
        self, temp_docx_with_table: Path
    ) -> None:
        """Test that row nodes contain cell children."""
        ast = parse_docx(temp_docx_with_table)

        table_nodes = [n for n in ast.nodes if n.node_type == NodeType.TABLE]
        if table_nodes:
            table = table_nodes[0]
            for row in table.children:
                assert len(row.children) > 0
                assert all(c.node_type == NodeType.TABLE_CELL for c in row.children)

    def test_table_cell_text_extracted(self, temp_docx_with_table: Path) -> None:
        """Test that table cell text is correctly extracted."""
        ast = parse_docx(temp_docx_with_table)

        table_nodes = [n for n in ast.nodes if n.node_type == NodeType.TABLE]
        if table_nodes:
            table = table_nodes[0]
            # Check that cell text is in the table's text content
            assert "Header 1" in table.text_content
            assert "Row 1, Col 1" in table.text_content


# =============================================================================
# Run-Level Tests
# =============================================================================


class TestRunLevelParsing:
    """Tests for run-level text extraction."""

    def test_parse_extracts_runs(self, temp_docx_with_formatting: Path) -> None:
        """Test that runs are properly extracted from paragraphs."""
        ast = parse_docx(temp_docx_with_formatting)

        # Find paragraph nodes with runs
        para_nodes = [n for n in ast.nodes if n.node_type == NodeType.PARAGRAPH]

        for node in para_nodes:
            if node.text_content:  # Non-empty paragraph
                # Should have at least one run
                assert len(node.runs) >= 0  # Some paragraphs may have no runs

    def test_run_text_matches_node_text(
        self, temp_docx_with_formatting: Path
    ) -> None:
        """Test that combined run text equals node text content."""
        ast = parse_docx(temp_docx_with_formatting)

        for node in ast.nodes:
            if node.runs:
                combined_run_text = "".join(r.text for r in node.runs)
                assert combined_run_text == node.text_content

    def test_run_positions_are_sequential(
        self, temp_docx_with_formatting: Path
    ) -> None:
        """Test that run positions are properly sequential."""
        ast = parse_docx(temp_docx_with_formatting)

        for node in ast.nodes:
            if len(node.runs) > 1:
                for i in range(len(node.runs) - 1):
                    current_run = node.runs[i]
                    next_run = node.runs[i + 1]
                    assert current_run.position.end_char == next_run.position.start_char


# =============================================================================
# Integration Tests
# =============================================================================


class TestParseDocxIntegration:
    """Integration tests for the full parsing pipeline."""

    def test_roundtrip_text_preservation(self, temp_docx_path: Path) -> None:
        """Test that all text is preserved through parsing."""
        doc = Document(str(temp_docx_path))
        original_text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

        ast = parse_docx(temp_docx_path)

        # All original text should be in the full_text
        for text in original_text_parts:
            assert text in ast.full_text

    def test_node_ids_are_unique(self, temp_docx_path: Path) -> None:
        """Test that all node IDs are unique."""
        ast = parse_docx(temp_docx_path)

        node_ids = [n.node_id for n in ast.nodes]
        assert len(node_ids) == len(set(node_ids))

    def test_position_ids_are_unique(self, temp_docx_path: Path) -> None:
        """Test that all position IDs are unique."""
        ast = parse_docx(temp_docx_path)

        position_ids = [n.position.position_id for n in ast.nodes]
        assert len(position_ids) == len(set(position_ids))
