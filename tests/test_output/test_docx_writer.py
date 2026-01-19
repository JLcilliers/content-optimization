"""
Tests for DOCX Writer Module

Tests the writing of DOCX documents with highlighting, including:
- Writing optimized documents
- Writing from AST
- FAQ section insertion
- Content insertion at positions
- Document validation
- Document merging
"""

from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from seo_optimizer.diffing.models import (
    Addition,
    ChangeSet,
    ChangeType,
    HighlightRegion,
    OptimizedContent,
    ParagraphContent,
    TextSegment,
)
from seo_optimizer.ingestion.models import (
    ContentNode,
    DocumentAST,
    DocumentMetadata,
    NodeType,
    PositionInfo,
)
from seo_optimizer.output.docx_writer import (
    OptimizedDocumentWriter,
    _add_paragraph_with_highlighting,
    add_color_legend,
    add_faq_section_enhanced,
    add_optimization_header,
    apply_heading_style,
    insert_content_at_position,
    insert_faq_section,
    merge_documents,
    setup_document_styles,
    validate_output,
    write_document_from_ast,
    write_optimized_docx,
)


# =============================================================================
# Fixtures
# =============================================================================


@pytest.fixture
def temp_source_docx(tmp_path: Path) -> Path:
    """Create a temporary source DOCX file."""
    doc = Document()
    doc.add_heading("Original Document", level=1)
    doc.add_paragraph("This is the original content.")
    doc.add_paragraph("Another paragraph in the original.")

    docx_path = tmp_path / "source.docx"
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def simple_ast() -> DocumentAST:
    """Create a simple DocumentAST for testing."""
    position = PositionInfo(
        position_id="p0",
        start_char=0,
        end_char=11,
    )
    node = ContentNode(
        node_id="node_p0",
        node_type=NodeType.PARAGRAPH,
        position=position,
        text_content="Hello world",
    )
    return DocumentAST(
        doc_id="test_doc",
        nodes=[node],
        metadata=DocumentMetadata(),
        full_text="Hello world",
        char_count=11,
    )


@pytest.fixture
def ast_with_heading() -> DocumentAST:
    """Create a DocumentAST with heading and paragraphs."""
    heading_node = ContentNode(
        node_id="node_h0",
        node_type=NodeType.HEADING,
        position=PositionInfo(position_id="h0", start_char=0, end_char=10),
        text_content="Test Title",
        metadata={"heading_level": 1},
    )
    para_node = ContentNode(
        node_id="node_p0",
        node_type=NodeType.PARAGRAPH,
        position=PositionInfo(position_id="p0", start_char=11, end_char=30),
        text_content="This is content.",
    )
    return DocumentAST(
        doc_id="test_doc",
        nodes=[heading_node, para_node],
        metadata=DocumentMetadata(),
        full_text="Test Title\nThis is content.",
        char_count=27,
    )


@pytest.fixture
def sample_changeset() -> ChangeSet:
    """Create a sample ChangeSet for testing."""
    region = HighlightRegion(
        node_id="node_p0",
        start_char=0,
        end_char=5,
        text="Hello",
        confidence=0.95,
    )
    addition = Addition(
        addition_id="add_001",
        node_ids=["node_p0"],
        highlight_regions=[region],
        total_text="Hello",
        confidence=0.95,
    )
    return ChangeSet(
        changeset_id="cs_001",
        original_doc_id="test_doc",
        optimized_doc_id="test_doc_opt",
        additions=[addition],
    )


@pytest.fixture
def empty_changeset() -> ChangeSet:
    """Create an empty ChangeSet."""
    return ChangeSet(
        changeset_id="cs_empty",
        original_doc_id="test_doc",
        optimized_doc_id="test_doc_opt",
        additions=[],
    )


# =============================================================================
# Add Paragraph with Highlighting Tests
# =============================================================================


class TestAddParagraphWithHighlighting:
    """Tests for _add_paragraph_with_highlighting function."""

    def test_add_paragraph_basic(self) -> None:
        """Test adding a basic paragraph."""
        doc = Document()

        para = _add_paragraph_with_highlighting(doc, "Test text")

        assert para is not None
        assert "Test text" in para.text

    def test_add_paragraph_with_style(self) -> None:
        """Test adding a paragraph with style."""
        doc = Document()

        para = _add_paragraph_with_highlighting(doc, "Heading", style="Heading 1")

        assert para is not None
        # Style should be applied
        assert para.style is not None

    def test_add_paragraph_with_highlight(self) -> None:
        """Test adding a paragraph with highlight."""
        doc = Document()

        para = _add_paragraph_with_highlighting(doc, "New text", highlight=True)

        assert para is not None
        # Check that highlight is applied (BRIGHT_GREEN for new content)
        for run in para.runs:
            assert run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN

    def test_add_paragraph_invalid_style_handled(self) -> None:
        """Test that invalid style is handled gracefully."""
        doc = Document()

        # Should not raise, even with invalid style
        para = _add_paragraph_with_highlighting(
            doc, "Test", style="NonexistentStyle123"
        )

        assert para is not None


# =============================================================================
# Write Optimized DOCX Tests
# =============================================================================


class TestWriteOptimizedDocx:
    """Tests for write_optimized_docx function."""

    def test_write_creates_output_file(
        self,
        tmp_path: Path,
        temp_source_docx: Path,
        simple_ast: DocumentAST,
        empty_changeset: ChangeSet,
    ) -> None:
        """Test that write creates an output file."""
        output_path = tmp_path / "output.docx"

        result = write_optimized_docx(
            temp_source_docx, simple_ast, empty_changeset, output_path
        )

        assert result == output_path
        assert output_path.exists()

    def test_write_preserves_original_content(
        self,
        tmp_path: Path,
        temp_source_docx: Path,
        simple_ast: DocumentAST,
        empty_changeset: ChangeSet,
    ) -> None:
        """Test that original content is preserved."""
        output_path = tmp_path / "output.docx"

        write_optimized_docx(temp_source_docx, simple_ast, empty_changeset, output_path)

        # Load and check content
        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Original Document" in full_text
        assert "original content" in full_text

    def test_write_missing_original_raises(
        self,
        tmp_path: Path,
        simple_ast: DocumentAST,
        empty_changeset: ChangeSet,
    ) -> None:
        """Test that missing original file raises error."""
        nonexistent = tmp_path / "nonexistent.docx"
        output_path = tmp_path / "output.docx"

        with pytest.raises(FileNotFoundError):
            write_optimized_docx(
                nonexistent, simple_ast, empty_changeset, output_path
            )


# =============================================================================
# Write Document from AST Tests
# =============================================================================


class TestWriteDocumentFromAst:
    """Tests for write_document_from_ast function."""

    def test_write_creates_file(self, tmp_path: Path, simple_ast: DocumentAST) -> None:
        """Test that write creates a file."""
        output_path = tmp_path / "ast_output.docx"

        result = write_document_from_ast(simple_ast, output_path)

        assert result == output_path
        assert output_path.exists()

    def test_write_includes_content(
        self, tmp_path: Path, simple_ast: DocumentAST
    ) -> None:
        """Test that AST content is written."""
        output_path = tmp_path / "ast_output.docx"

        write_document_from_ast(simple_ast, output_path)

        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Hello world" in full_text

    def test_write_with_heading(
        self, tmp_path: Path, ast_with_heading: DocumentAST
    ) -> None:
        """Test writing AST with heading."""
        output_path = tmp_path / "heading_output.docx"

        write_document_from_ast(ast_with_heading, output_path)

        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Test Title" in full_text
        assert "This is content" in full_text

    def test_write_highlights_new_nodes(
        self, tmp_path: Path, simple_ast: DocumentAST
    ) -> None:
        """Test that new nodes are highlighted."""
        output_path = tmp_path / "highlighted_output.docx"
        new_node_ids = {"node_p0"}  # Mark the node as new

        write_document_from_ast(
            simple_ast, output_path, highlight_new=True, new_node_ids=new_node_ids
        )

        doc = Document(str(output_path))
        # Check that content is highlighted (BRIGHT_GREEN for new content)
        for para in doc.paragraphs:
            if "Hello world" in para.text:
                for run in para.runs:
                    assert run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN


# =============================================================================
# Insert FAQ Section Tests
# =============================================================================


class TestInsertFaqSection:
    """Tests for insert_faq_section function."""

    def test_insert_faq_basic(self) -> None:
        """Test inserting a basic FAQ section."""
        doc = Document()
        doc.add_paragraph("Existing content")

        faq_content = [
            ("What is this?", "This is a test product."),
            ("How does it work?", "It works automatically."),
        ]

        result = insert_faq_section(doc, faq_content)

        assert result is doc
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Frequently Asked Questions" in full_text
        assert "What is this?" in full_text
        assert "This is a test product" in full_text

    def test_insert_faq_with_highlight(self) -> None:
        """Test that FAQ is highlighted when requested."""
        doc = Document()

        faq_content = [("Question?", "Answer.")]
        insert_faq_section(doc, faq_content, highlight=True)

        # Check for highlighted content (green for new content)
        highlighted = False
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN:
                    highlighted = True
                    break
        assert highlighted

    def test_insert_faq_without_highlight(self) -> None:
        """Test that FAQ is not highlighted when not requested."""
        doc = Document()

        faq_content = [("Question?", "Answer.")]
        insert_faq_section(doc, faq_content, highlight=False)

        # Check that no content is highlighted
        for para in doc.paragraphs:
            for run in para.runs:
                assert run.font.highlight_color is None

    def test_insert_multiple_faqs(self) -> None:
        """Test inserting multiple FAQ items."""
        doc = Document()

        faq_content = [
            ("Q1?", "A1"),
            ("Q2?", "A2"),
            ("Q3?", "A3"),
        ]
        insert_faq_section(doc, faq_content)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Q1?" in full_text
        assert "Q2?" in full_text
        assert "Q3?" in full_text


# =============================================================================
# Insert Content at Position Tests
# =============================================================================


class TestInsertContentAtPosition:
    """Tests for insert_content_at_position function."""

    def test_insert_at_end(self, temp_source_docx: Path) -> None:
        """Test inserting content at the end."""
        doc = Document(str(temp_source_docx))
        initial_count = len(doc.paragraphs)

        para = insert_content_at_position(doc, "New content", position=initial_count)

        assert para is not None
        assert "New content" in para.text

    def test_insert_with_highlight(self, temp_source_docx: Path) -> None:
        """Test that inserted content is highlighted."""
        doc = Document(str(temp_source_docx))

        para = insert_content_at_position(doc, "New content", position=0, highlight=True)

        for run in para.runs:
            assert run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN

    def test_insert_without_highlight(self, temp_source_docx: Path) -> None:
        """Test that content is not highlighted when not requested."""
        doc = Document(str(temp_source_docx))

        para = insert_content_at_position(
            doc, "New content", position=0, highlight=False
        )

        for run in para.runs:
            assert run.font.highlight_color is None


# =============================================================================
# Validate Output Tests
# =============================================================================


class TestValidateOutput:
    """Tests for validate_output function."""

    def test_validate_existing_docx(self, temp_source_docx: Path) -> None:
        """Test validating an existing DOCX file."""
        warnings = validate_output(temp_source_docx)

        # Should return list (may have warnings about no highlights)
        assert isinstance(warnings, list)

    def test_validate_nonexistent_file(self, tmp_path: Path) -> None:
        """Test validating a nonexistent file."""
        nonexistent = tmp_path / "nonexistent.docx"

        warnings = validate_output(nonexistent)

        assert "File does not exist" in warnings

    def test_validate_wrong_extension(self, tmp_path: Path) -> None:
        """Test validating file with wrong extension."""
        txt_file = tmp_path / "document.txt"
        txt_file.write_text("Hello")

        warnings = validate_output(txt_file)

        assert any("extension" in w.lower() for w in warnings)

    def test_validate_invalid_docx(self, tmp_path: Path) -> None:
        """Test validating an invalid DOCX file."""
        invalid = tmp_path / "invalid.docx"
        invalid.write_text("This is not a DOCX")

        warnings = validate_output(invalid)

        assert any("Invalid" in w for w in warnings)

    def test_validate_with_highlights(self, tmp_path: Path) -> None:
        """Test validation of document with highlights."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Highlighted text")
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        output_path = tmp_path / "highlighted.docx"
        doc.save(str(output_path))

        warnings = validate_output(output_path)

        # Should not warn about no highlights
        assert not any("No highlighted" in w for w in warnings)


# =============================================================================
# Merge Documents Tests
# =============================================================================


class TestMergeDocuments:
    """Tests for merge_documents function."""

    def test_merge_appends_content(
        self, tmp_path: Path, temp_source_docx: Path
    ) -> None:
        """Test merging appends content to end."""
        output_path = tmp_path / "merged.docx"
        additions = [("end", "New content at end")]

        result = merge_documents(temp_source_docx, additions, output_path)

        assert result == output_path
        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "New content at end" in full_text

    def test_merge_prepends_content(
        self, tmp_path: Path, temp_source_docx: Path
    ) -> None:
        """Test merging prepends content to start."""
        output_path = tmp_path / "merged.docx"
        additions = [("start", "New content at start")]

        merge_documents(temp_source_docx, additions, output_path)

        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "New content at start" in full_text

    def test_merge_with_highlight(
        self, tmp_path: Path, temp_source_docx: Path
    ) -> None:
        """Test that merged content is highlighted."""
        output_path = tmp_path / "merged.docx"
        additions = [("end", "Highlighted addition")]

        merge_documents(
            temp_source_docx, additions, output_path, highlight_additions=True
        )

        doc = Document(str(output_path))
        # Find the paragraph with the addition (BRIGHT_GREEN for new content)
        for para in doc.paragraphs:
            if "Highlighted addition" in para.text:
                for run in para.runs:
                    assert run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN

    def test_merge_without_highlight(
        self, tmp_path: Path, temp_source_docx: Path
    ) -> None:
        """Test merging without highlighting."""
        output_path = tmp_path / "merged.docx"
        additions = [("end", "Non-highlighted addition")]

        merge_documents(
            temp_source_docx, additions, output_path, highlight_additions=False
        )

        doc = Document(str(output_path))
        for para in doc.paragraphs:
            if "Non-highlighted addition" in para.text:
                for run in para.runs:
                    assert run.font.highlight_color is None

    def test_merge_multiple_additions(
        self, tmp_path: Path, temp_source_docx: Path
    ) -> None:
        """Test merging multiple additions."""
        output_path = tmp_path / "merged.docx"
        additions = [
            ("end", "First addition"),
            ("end", "Second addition"),
        ]

        merge_documents(temp_source_docx, additions, output_path)

        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "First addition" in full_text
        assert "Second addition" in full_text


# =============================================================================
# Integration Tests
# =============================================================================


class TestDocxWriterIntegration:
    """Integration tests for the DOCX writer module."""

    def test_full_workflow(
        self,
        tmp_path: Path,
        temp_source_docx: Path,
        ast_with_heading: DocumentAST,
        sample_changeset: ChangeSet,
    ) -> None:
        """Test the full write workflow."""
        output_path = tmp_path / "full_workflow.docx"

        # Write optimized document
        result = write_optimized_docx(
            temp_source_docx, ast_with_heading, sample_changeset, output_path
        )

        # Validate output
        warnings = validate_output(result)

        # Should be a valid DOCX
        assert "Invalid" not in str(warnings)
        assert output_path.exists()

    def test_write_then_validate(
        self, tmp_path: Path, simple_ast: DocumentAST
    ) -> None:
        """Test writing and then validating."""
        output_path = tmp_path / "write_validate.docx"

        write_document_from_ast(simple_ast, output_path)
        warnings = validate_output(output_path)

        # Should be valid
        assert "Invalid" not in str(warnings)
        assert "does not exist" not in str(warnings)


# =============================================================================
# Heading Style Tests
# =============================================================================


class TestApplyHeadingStyle:
    """Tests for apply_heading_style function."""

    def test_parses_h1_marker(self) -> None:
        """Test parsing [H1] marker."""
        doc = Document()
        para = doc.add_paragraph()

        result = apply_heading_style(para, "[H1] Main Title")

        assert result == "Main Title"

    def test_parses_h2_marker(self) -> None:
        """Test parsing [H2] marker."""
        doc = Document()
        para = doc.add_paragraph()

        result = apply_heading_style(para, "[H2] Section Header")

        assert result == "Section Header"

    def test_parses_h3_marker(self) -> None:
        """Test parsing [H3] marker."""
        doc = Document()
        para = doc.add_paragraph()

        result = apply_heading_style(para, "[H3] Subsection")

        assert result == "Subsection"

    def test_no_marker_returns_original(self) -> None:
        """Test that text without marker is returned unchanged."""
        doc = Document()
        para = doc.add_paragraph()

        result = apply_heading_style(para, "Regular paragraph text")

        assert result == "Regular paragraph text"

    def test_invalid_marker_not_parsed(self) -> None:
        """Test that invalid markers are not parsed."""
        doc = Document()
        para = doc.add_paragraph()

        result = apply_heading_style(para, "[H4] Level 4 heading")

        assert result == "[H4] Level 4 heading"


class TestSetupDocumentStyles:
    """Tests for setup_document_styles function."""

    def test_styles_are_configured(self) -> None:
        """Test that document styles are configured."""
        doc = Document()

        setup_document_styles(doc)

        # Should not raise and styles should exist
        assert "Heading 1" in [s.name for s in doc.styles]
        assert "Heading 2" in [s.name for s in doc.styles]
        assert "Normal" in [s.name for s in doc.styles]


# =============================================================================
# Document Header Tests
# =============================================================================


class TestAddOptimizationHeader:
    """Tests for add_optimization_header function."""

    def test_adds_title(self) -> None:
        """Test that header adds title."""
        doc = Document()

        add_optimization_header(doc, url="https://example.com", keyword="test")

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "SEO OPTIMIZATION REPORT" in full_text

    def test_adds_metadata_table(self) -> None:
        """Test that header adds metadata table."""
        doc = Document()

        add_optimization_header(doc, url="https://example.com", keyword="test keyword")

        # Should have a table
        assert len(doc.tables) > 0
        table = doc.tables[0]
        table_text = " ".join(cell.text for row in table.rows for cell in row.cells)
        assert "URL:" in table_text
        assert "Target Keyword:" in table_text

    def test_uses_content_object(self) -> None:
        """Test header with OptimizedContent object."""
        doc = Document()
        content = OptimizedContent(
            url="https://test.com",
            target_keyword="SEO",
            optimization_date="2024-01-15",
            change_summary={"total": 10},
        )

        add_optimization_header(doc, content=content)

        # Should use values from content
        table = doc.tables[0]
        table_text = " ".join(cell.text for row in table.rows for cell in row.cells)
        assert "https://test.com" in table_text
        assert "SEO" in table_text


class TestAddColorLegend:
    """Tests for add_color_legend function."""

    def test_adds_legend(self) -> None:
        """Test that color legend is added."""
        doc = Document()

        para = add_color_legend(doc)

        assert "CHANGE LEGEND" in para.text
        assert "New Content" in para.text
        assert "Modified" in para.text
        assert "Removed" in para.text

    def test_legend_has_highlighted_samples(self) -> None:
        """Test that legend contains highlighted samples."""
        doc = Document()

        para = add_color_legend(doc)

        # Check for runs with highlight colors
        has_green = False
        has_yellow = False
        has_strike = False
        for run in para.runs:
            if run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN:
                has_green = True
            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                has_yellow = True
            if run.font.strike:
                has_strike = True

        assert has_green, "Legend should have green highlighted sample"
        assert has_yellow, "Legend should have yellow highlighted sample"
        assert has_strike, "Legend should have strikethrough sample"


# =============================================================================
# Enhanced FAQ Section Tests
# =============================================================================


class TestAddFaqSectionEnhanced:
    """Tests for add_faq_section_enhanced function."""

    def test_adds_section_header(self) -> None:
        """Test that FAQ section header is added."""
        doc = Document()
        faqs = [{"question": "What is this?", "answer": "This is a test."}]

        add_faq_section_enhanced(doc, faqs)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Frequently Asked Questions" in full_text

    def test_adds_questions_and_answers(self) -> None:
        """Test that questions and answers are added."""
        doc = Document()
        faqs = [
            {"question": "Q1?", "answer": "A1"},
            {"question": "Q2?", "answer": "A2"},
        ]

        add_faq_section_enhanced(doc, faqs)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Q1?" in full_text
        assert "A1" in full_text
        assert "Q2?" in full_text
        assert "A2" in full_text

    def test_highlights_faq_content(self) -> None:
        """Test that FAQ content is highlighted green."""
        doc = Document()
        faqs = [{"question": "Test?", "answer": "Answer"}]

        add_faq_section_enhanced(doc, faqs, highlight=True)

        # Find highlighted content
        has_highlight = False
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN:
                    has_highlight = True
                    break

        assert has_highlight

    def test_no_highlight_when_disabled(self) -> None:
        """Test that FAQ content is not highlighted when disabled."""
        doc = Document()
        faqs = [{"question": "Test?", "answer": "Answer"}]

        add_faq_section_enhanced(doc, faqs, highlight=False)

        # Check no highlighting
        for para in doc.paragraphs:
            for run in para.runs:
                assert run.font.highlight_color is None

    def test_empty_faqs_does_nothing(self) -> None:
        """Test that empty FAQ list does nothing."""
        doc = Document()
        initial_count = len(doc.paragraphs)

        add_faq_section_enhanced(doc, [])

        assert len(doc.paragraphs) == initial_count


# =============================================================================
# Optimized Document Writer Tests
# =============================================================================


class TestOptimizedDocumentWriter:
    """Tests for OptimizedDocumentWriter class."""

    def test_create_document_basic(self) -> None:
        """Test creating a basic document."""
        writer = OptimizedDocumentWriter()
        content = OptimizedContent(
            url="https://example.com",
            target_keyword="test",
        )

        doc = writer.create_document(content)

        assert doc is not None
        assert len(doc.paragraphs) > 0

    def test_create_document_with_meta(self) -> None:
        """Test creating document with meta information."""
        writer = OptimizedDocumentWriter()
        content = OptimizedContent(
            url="https://example.com",
            target_keyword="test",
            meta_title=ParagraphContent(
                segments=[TextSegment(text="Page Title", change_type=ChangeType.UNCHANGED)]
            ),
            meta_description=ParagraphContent(
                segments=[TextSegment(text="Description", change_type=ChangeType.UNCHANGED)]
            ),
        )

        doc = writer.create_document(content)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Meta Information" in full_text
        assert "Page Title" in full_text

    def test_create_document_with_body(self) -> None:
        """Test creating document with body content."""
        writer = OptimizedDocumentWriter()
        content = OptimizedContent(
            url="https://example.com",
            target_keyword="test",
            body_paragraphs=[
                ParagraphContent(
                    segments=[TextSegment(text="Body content", change_type=ChangeType.UNCHANGED)]
                )
            ],
        )

        doc = writer.create_document(content)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Body content" in full_text

    def test_create_document_with_faq(self) -> None:
        """Test creating document with FAQ section."""
        writer = OptimizedDocumentWriter()
        content = OptimizedContent(
            url="https://example.com",
            target_keyword="test",
            faq_items=[{"question": "Q?", "answer": "A."}],
        )

        doc = writer.create_document(content)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Frequently Asked Questions" in full_text
        assert "Q?" in full_text

    def test_write_to_file(self, tmp_path: Path) -> None:
        """Test writing document to file."""
        writer = OptimizedDocumentWriter()
        content = OptimizedContent(
            url="https://example.com",
            target_keyword="test",
        )
        output_path = tmp_path / "output.docx"

        result = writer.write_to_file(content, output_path)

        assert result == output_path
        assert output_path.exists()

    def test_write_to_stream(self) -> None:
        """Test writing document to stream."""
        import io

        writer = OptimizedDocumentWriter()
        content = OptimizedContent(
            url="https://example.com",
            target_keyword="test",
        )
        stream = io.BytesIO()

        writer.write_to_stream(content, stream)

        assert stream.tell() == 0  # Should be seeked back to start
        assert len(stream.getvalue()) > 0  # Should have content

    def test_change_highlighting_applied(self) -> None:
        """Test that change highlighting is applied."""
        writer = OptimizedDocumentWriter()
        content = OptimizedContent(
            url="https://example.com",
            target_keyword="test",
            body_paragraphs=[
                ParagraphContent(
                    segments=[
                        TextSegment(text="Original ", change_type=ChangeType.UNCHANGED),
                        TextSegment(text="new content", change_type=ChangeType.INSERTED),
                    ]
                )
            ],
        )

        doc = writer.create_document(content)

        # Find the highlighted content
        has_green = False
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN:
                    has_green = True
                    break

        assert has_green
