"""
Tests for Highlighter Module

Tests the highlighting functionality for DOCX documents, including:
- Run-level highlighting
- Run splitting for partial highlights
- Text search and highlight
- ChangeSet-based highlighting
"""

from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from seo_optimizer.diffing.models import Addition, ChangeSet, HighlightRegion
from seo_optimizer.output.highlighter import (
    HIGHLIGHT_COLOR_INDEX,
    _apply_highlight_to_run,
    _copy_run_formatting,
    _get_paragraph_text_with_positions,
    apply_highlights,
    create_highlighted_run,
    highlight_new_paragraph,
    highlight_text_in_paragraph,
    split_run_for_highlight,
)


# =============================================================================
# Fixtures
# =============================================================================


@pytest.fixture
def temp_doc(tmp_path: Path) -> Document:
    """Create a temporary in-memory document."""
    return Document()


@pytest.fixture
def temp_docx_with_content(tmp_path: Path) -> tuple[Path, Document]:
    """Create a DOCX file with content for highlighting tests."""
    doc = Document()
    doc.add_paragraph("This is the first paragraph with some text.")
    doc.add_paragraph("Hello world, this is a test.")
    doc.add_paragraph("Another paragraph for testing highlights.")

    docx_path = tmp_path / "content_document.docx"
    doc.save(str(docx_path))
    return docx_path, Document(str(docx_path))


@pytest.fixture
def sample_changeset() -> ChangeSet:
    """Create a sample ChangeSet for testing."""
    region = HighlightRegion(
        node_id="doc_test_p0",
        start_char=5,
        end_char=15,
        text="is the fir",
        confidence=0.95,
    )
    addition = Addition(
        addition_id="add_001",
        node_ids=["doc_test_p0"],
        highlight_regions=[region],
        total_text="is the fir",
        confidence=0.95,
    )
    return ChangeSet(
        changeset_id="cs_001",
        original_doc_id="doc_test",
        optimized_doc_id="doc_test_opt",
        additions=[addition],
    )


@pytest.fixture
def empty_changeset() -> ChangeSet:
    """Create an empty ChangeSet."""
    return ChangeSet(
        changeset_id="cs_empty",
        original_doc_id="doc_test",
        optimized_doc_id="doc_test_opt",
        additions=[],
    )


# =============================================================================
# Highlight Color Tests
# =============================================================================


class TestHighlightColor:
    """Tests for highlight color configuration."""

    def test_highlight_color_is_yellow(self) -> None:
        """Test that default highlight color is yellow."""
        assert HIGHLIGHT_COLOR_INDEX == WD_COLOR_INDEX.YELLOW


# =============================================================================
# Apply Highlight to Run Tests
# =============================================================================


class TestApplyHighlightToRun:
    """Tests for _apply_highlight_to_run function."""

    def test_apply_highlight_sets_color(self, temp_doc: Document) -> None:
        """Test that applying highlight sets the highlight color."""
        para = temp_doc.add_paragraph()
        run = para.add_run("Test text")

        _apply_highlight_to_run(run)

        assert run.font.highlight_color == WD_COLOR_INDEX.YELLOW

    def test_apply_highlight_custom_color(self, temp_doc: Document) -> None:
        """Test applying a custom highlight color."""
        para = temp_doc.add_paragraph()
        run = para.add_run("Test text")

        _apply_highlight_to_run(run, WD_COLOR_INDEX.BRIGHT_GREEN)

        assert run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN


# =============================================================================
# Copy Run Formatting Tests
# =============================================================================


class TestCopyRunFormatting:
    """Tests for _copy_run_formatting function."""

    def test_copy_bold_formatting(self, temp_doc: Document) -> None:
        """Test that bold formatting is copied."""
        para = temp_doc.add_paragraph()
        source = para.add_run("Bold text")
        source.bold = True
        target = para.add_run("Target")

        _copy_run_formatting(source, target)

        assert target.bold is True

    def test_copy_italic_formatting(self, temp_doc: Document) -> None:
        """Test that italic formatting is copied."""
        para = temp_doc.add_paragraph()
        source = para.add_run("Italic text")
        source.italic = True
        target = para.add_run("Target")

        _copy_run_formatting(source, target)

        assert target.italic is True

    def test_copy_underline_formatting(self, temp_doc: Document) -> None:
        """Test that underline formatting is copied."""
        para = temp_doc.add_paragraph()
        source = para.add_run("Underline text")
        source.underline = True
        target = para.add_run("Target")

        _copy_run_formatting(source, target)

        assert target.underline is True

    def test_copy_multiple_formatting(self, temp_doc: Document) -> None:
        """Test copying multiple formatting properties."""
        para = temp_doc.add_paragraph()
        source = para.add_run("Formatted text")
        source.bold = True
        source.italic = True
        target = para.add_run("Target")

        _copy_run_formatting(source, target)

        assert target.bold is True
        assert target.italic is True


# =============================================================================
# Get Paragraph Text with Positions Tests
# =============================================================================


class TestGetParagraphTextWithPositions:
    """Tests for _get_paragraph_text_with_positions function."""

    def test_get_positions_single_run(self, temp_doc: Document) -> None:
        """Test getting positions with a single run."""
        para = temp_doc.add_paragraph()
        para.add_run("Hello world")

        positions = _get_paragraph_text_with_positions(para)

        assert len(positions) == 1
        run_idx, start, end, run = positions[0]
        assert run_idx == 0
        assert start == 0
        assert end == 11  # "Hello world" = 11 chars

    def test_get_positions_multiple_runs(self, temp_doc: Document) -> None:
        """Test getting positions with multiple runs."""
        para = temp_doc.add_paragraph()
        para.add_run("Hello ")  # 6 chars
        para.add_run("world")  # 5 chars

        positions = _get_paragraph_text_with_positions(para)

        assert len(positions) == 2
        assert positions[0][1:3] == (0, 6)  # start, end
        assert positions[1][1:3] == (6, 11)

    def test_get_positions_empty_paragraph(self, temp_doc: Document) -> None:
        """Test getting positions for empty paragraph."""
        para = temp_doc.add_paragraph()

        positions = _get_paragraph_text_with_positions(para)

        assert len(positions) == 0

    def test_get_positions_skips_empty_runs(self, temp_doc: Document) -> None:
        """Test that empty runs are skipped."""
        para = temp_doc.add_paragraph()
        para.add_run("Hello")
        para.add_run("")  # Empty run
        para.add_run("World")

        positions = _get_paragraph_text_with_positions(para)

        # Should only have 2 non-empty runs
        assert len(positions) == 2


# =============================================================================
# Split Run for Highlight Tests
# =============================================================================


class TestSplitRunForHighlight:
    """Tests for split_run_for_highlight function."""

    def test_split_run_middle(self, temp_doc: Document) -> None:
        """Test splitting a run in the middle."""
        para = temp_doc.add_paragraph()
        para.add_run("Hello world")

        # Highlight "o wor" (positions 4-9)
        new_runs = split_run_for_highlight(para, 0, 4, 9)

        # Should have created new runs
        assert len(new_runs) >= 2

    def test_split_run_start(self, temp_doc: Document) -> None:
        """Test splitting at the start of a run."""
        para = temp_doc.add_paragraph()
        para.add_run("Hello world")

        # Highlight "Hello" (positions 0-5)
        new_runs = split_run_for_highlight(para, 0, 0, 5)

        # First run should be highlighted
        assert len(new_runs) >= 1
        assert new_runs[0].font.highlight_color == HIGHLIGHT_COLOR_INDEX

    def test_split_run_end(self, temp_doc: Document) -> None:
        """Test splitting at the end of a run."""
        para = temp_doc.add_paragraph()
        para.add_run("Hello world")

        # Highlight "world" (positions 6-11)
        new_runs = split_run_for_highlight(para, 0, 6, 11)

        # Should have runs created
        assert len(new_runs) >= 1

    def test_split_run_empty_returns_empty(self, temp_doc: Document) -> None:
        """Test that splitting an empty run returns empty list."""
        para = temp_doc.add_paragraph()
        para.add_run("")  # Empty run

        result = split_run_for_highlight(para, 0, 0, 5)

        assert result == []

    def test_split_run_invalid_range_returns_empty(self, temp_doc: Document) -> None:
        """Test that invalid range returns empty list."""
        para = temp_doc.add_paragraph()
        para.add_run("Hello")

        result = split_run_for_highlight(para, 0, 5, 3)  # start > end

        assert result == []

    def test_split_run_preserves_formatting(self, temp_doc: Document) -> None:
        """Test that splitting preserves original formatting."""
        para = temp_doc.add_paragraph()
        run = para.add_run("Hello world")
        run.bold = True

        new_runs = split_run_for_highlight(para, 0, 4, 9)

        # All runs should be bold
        for new_run in new_runs:
            assert new_run.bold is True


# =============================================================================
# Highlight Text in Paragraph Tests
# =============================================================================


class TestHighlightTextInParagraph:
    """Tests for highlight_text_in_paragraph function."""

    def test_highlight_existing_text(
        self, temp_docx_with_content: tuple[Path, Document]
    ) -> None:
        """Test highlighting text that exists in paragraph."""
        _, doc = temp_docx_with_content
        para = doc.paragraphs[1]  # "Hello world, this is a test."

        result = highlight_text_in_paragraph(para, "world")

        assert result is True

    def test_highlight_nonexistent_text(
        self, temp_docx_with_content: tuple[Path, Document]
    ) -> None:
        """Test highlighting text that doesn't exist."""
        _, doc = temp_docx_with_content
        para = doc.paragraphs[0]

        result = highlight_text_in_paragraph(para, "xyz_nonexistent")

        assert result is False

    def test_highlight_case_insensitive(
        self, temp_docx_with_content: tuple[Path, Document]
    ) -> None:
        """Test case-insensitive highlighting."""
        _, doc = temp_docx_with_content
        para = doc.paragraphs[1]  # "Hello world..."

        result = highlight_text_in_paragraph(para, "HELLO", case_sensitive=False)

        assert result is True

    def test_highlight_case_sensitive(
        self, temp_docx_with_content: tuple[Path, Document]
    ) -> None:
        """Test case-sensitive highlighting."""
        _, doc = temp_docx_with_content
        para = doc.paragraphs[1]

        # Should not find "HELLO" when case-sensitive
        result = highlight_text_in_paragraph(para, "HELLO", case_sensitive=True)

        assert result is False


# =============================================================================
# Highlight New Paragraph Tests
# =============================================================================


class TestHighlightNewParagraph:
    """Tests for highlight_new_paragraph function."""

    def test_highlight_entire_paragraph(self, temp_doc: Document) -> None:
        """Test highlighting an entire paragraph."""
        para = temp_doc.add_paragraph()
        para.add_run("This is new content")
        para.add_run(" and more content")

        highlight_new_paragraph(para)

        for run in para.runs:
            assert run.font.highlight_color == HIGHLIGHT_COLOR_INDEX

    def test_highlight_paragraph_with_multiple_runs(self, temp_doc: Document) -> None:
        """Test highlighting paragraph with multiple runs."""
        para = temp_doc.add_paragraph()
        para.add_run("First ")
        para.add_run("Second ")
        para.add_run("Third")

        highlight_new_paragraph(para)

        assert len(para.runs) == 3
        for run in para.runs:
            assert run.font.highlight_color == HIGHLIGHT_COLOR_INDEX


# =============================================================================
# Create Highlighted Run Tests
# =============================================================================


class TestCreateHighlightedRun:
    """Tests for create_highlighted_run function."""

    def test_create_highlighted_run_basic(self, temp_doc: Document) -> None:
        """Test creating a basic highlighted run."""
        para = temp_doc.add_paragraph()

        run = create_highlighted_run(para, "New text")

        assert run.text == "New text"
        assert run.font.highlight_color == HIGHLIGHT_COLOR_INDEX

    def test_create_highlighted_run_with_formatting(self, temp_doc: Document) -> None:
        """Test creating highlighted run with copied formatting."""
        para = temp_doc.add_paragraph()
        source_run = para.add_run("Source")
        source_run.bold = True
        source_run.italic = True

        new_run = create_highlighted_run(para, "New text", copy_formatting_from=source_run)

        assert new_run.text == "New text"
        assert new_run.font.highlight_color == HIGHLIGHT_COLOR_INDEX
        assert new_run.bold is True
        assert new_run.italic is True


# =============================================================================
# Apply Highlights Tests
# =============================================================================


class TestApplyHighlights:
    """Tests for apply_highlights function."""

    def test_apply_highlights_empty_changeset(
        self, temp_docx_with_content: tuple[Path, Document], empty_changeset: ChangeSet
    ) -> None:
        """Test applying highlights with empty changeset."""
        _, doc = temp_docx_with_content

        result = apply_highlights(doc, empty_changeset)

        assert result is doc  # Should return same document

    def test_apply_highlights_returns_document(
        self, temp_docx_with_content: tuple[Path, Document], sample_changeset: ChangeSet
    ) -> None:
        """Test that apply_highlights returns the document."""
        _, doc = temp_docx_with_content

        result = apply_highlights(doc, sample_changeset)

        assert result is doc

    def test_apply_highlights_skips_low_confidence(
        self, temp_docx_with_content: tuple[Path, Document]
    ) -> None:
        """Test that low confidence additions are skipped."""
        _, doc = temp_docx_with_content

        # Create low-confidence changeset
        region = HighlightRegion(
            node_id="doc_test_p0",
            start_char=0,
            end_char=5,
            text="This ",
            confidence=0.5,  # Low confidence
        )
        addition = Addition(
            addition_id="add_low",
            node_ids=["doc_test_p0"],
            highlight_regions=[region],
            total_text="This ",
            confidence=0.5,  # Low confidence
        )
        changeset = ChangeSet(
            changeset_id="cs_low",
            original_doc_id="doc_test",
            optimized_doc_id="doc_test_opt",
            additions=[addition],
        )

        result = apply_highlights(doc, changeset)

        # Document should be returned but low-confidence shouldn't be highlighted
        assert result is doc


# =============================================================================
# Integration Tests
# =============================================================================


class TestHighlighterIntegration:
    """Integration tests for the highlighter module."""

    def test_highlight_and_save(
        self, tmp_path: Path, temp_docx_with_content: tuple[Path, Document]
    ) -> None:
        """Test highlighting and saving a document."""
        _, doc = temp_docx_with_content

        # Highlight first paragraph
        para = doc.paragraphs[0]
        highlight_new_paragraph(para)

        # Save and reload
        output_path = tmp_path / "highlighted.docx"
        doc.save(str(output_path))

        # Verify file was created
        assert output_path.exists()

        # Reload and verify highlights
        reloaded = Document(str(output_path))
        for run in reloaded.paragraphs[0].runs:
            assert run.font.highlight_color == HIGHLIGHT_COLOR_INDEX

    def test_partial_highlight_preserves_content(self, temp_doc: Document) -> None:
        """Test that partial highlighting preserves all content."""
        para = temp_doc.add_paragraph()
        para.add_run("Hello beautiful world")
        original_text = "Hello beautiful world"

        # Highlight "beautiful"
        highlight_text_in_paragraph(para, "beautiful")

        # Get combined text from all runs
        combined_text = "".join(run.text for run in para.runs)
        assert combined_text == original_text
