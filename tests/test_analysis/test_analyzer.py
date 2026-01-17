"""
Tests for Content Analyzer Module

Integration tests for the main orchestrator.
"""

from pathlib import Path

import pytest
from docx import Document

from seo_optimizer.analysis.analyzer import (
    AnalyzerConfig,
    ContentAnalyzer,
    analyze_content,
    analyze_docx,
)
from seo_optimizer.analysis.models import (
    AnalysisResult,
    GEOScore,
    KeywordConfig,
)
from seo_optimizer.ingestion.models import (
    ContentNode,
    DocumentAST,
    DocumentMetadata,
    NodeType,
    PositionInfo,
)


# Mark all tests that require model loading as slow
pytestmark = pytest.mark.slow


# =============================================================================
# Fixtures
# =============================================================================


@pytest.fixture
def analyzer() -> ContentAnalyzer:
    """Create a content analyzer instance."""
    return ContentAnalyzer()


@pytest.fixture
def simple_keywords() -> KeywordConfig:
    """Simple keyword configuration."""
    return KeywordConfig(
        primary_keyword="cloud computing",
        secondary_keywords=["AWS", "Azure", "serverless"],
        semantic_entities=["data center", "virtualization", "scalability"],
    )


@pytest.fixture
def temp_docx(tmp_path: Path) -> Path:
    """Create a temporary DOCX file for testing."""
    doc = Document()
    doc.add_heading("Cloud Computing Guide", level=1)
    doc.add_paragraph(
        "Cloud computing is the delivery of computing services over the internet. "
        "AWS and Azure are major providers. This guide covers serverless architecture."
    )
    doc.add_heading("What is Cloud Computing?", level=2)
    doc.add_paragraph(
        "Cloud computing allows businesses to use data centers and virtualization "
        "to achieve scalability without managing physical servers."
    )
    doc.add_heading("Benefits", level=2)
    doc.add_paragraph(
        "The main benefits include cost savings, flexibility, and improved performance."
    )
    doc.add_heading("Frequently Asked Questions", level=2)
    doc.add_paragraph("Q: What is AWS? A: Amazon Web Services is a cloud platform.")

    docx_path = tmp_path / "test_document.docx"
    doc.save(str(docx_path))
    return docx_path


def create_test_ast(
    full_text: str,
    nodes: list[ContentNode] | None = None,
) -> DocumentAST:
    """Helper to create test AST."""
    return DocumentAST(
        doc_id="test_doc",
        nodes=nodes or [],
        full_text=full_text,
        char_count=len(full_text),
        metadata=DocumentMetadata(),
    )


def create_heading_node(text: str, level: int, position: int = 0) -> ContentNode:
    """Helper to create heading node."""
    return ContentNode(
        node_id=f"h{level}_{position}",
        node_type=NodeType.HEADING,
        text_content=text,
        position=PositionInfo(
            position_id=f"pos_{position}",
            start_char=0,
            end_char=len(text),
        ),
        metadata={"level": level},
    )


def create_paragraph_node(text: str, position: int = 0) -> ContentNode:
    """Helper to create paragraph node."""
    return ContentNode(
        node_id=f"p_{position}",
        node_type=NodeType.PARAGRAPH,
        text_content=text,
        position=PositionInfo(
            position_id=f"pos_{position}",
            start_char=0,
            end_char=len(text),
        ),
    )


# =============================================================================
# Basic Analysis Tests
# =============================================================================


class TestBasicAnalysis:
    """Basic analysis functionality tests."""

    def test_analyze_returns_result(
        self, analyzer: ContentAnalyzer, simple_keywords: KeywordConfig
    ) -> None:
        """Test that analyze returns AnalysisResult."""
        nodes = [
            create_heading_node("Cloud Computing Overview", level=1),
            create_paragraph_node("Cloud computing enables scalable infrastructure."),
        ]
        ast = create_test_ast(
            "Cloud Computing Overview. Cloud computing enables scalable infrastructure.",
            nodes=nodes,
        )

        result = analyzer.analyze(ast, simple_keywords)

        assert isinstance(result, AnalysisResult)

    def test_analyze_has_geo_score(
        self, analyzer: ContentAnalyzer, simple_keywords: KeywordConfig
    ) -> None:
        """Test that result includes GEO score."""
        nodes = [
            create_heading_node("Cloud Computing", level=1),
            create_paragraph_node("Learn about cloud computing and AWS services."),
        ]
        ast = create_test_ast(
            "Cloud Computing. Learn about cloud computing and AWS services.",
            nodes=nodes,
        )

        result = analyzer.analyze(ast, simple_keywords)

        assert isinstance(result.geo_score, GEOScore)
        assert result.geo_score.total >= 0

    def test_analyze_has_document_stats(
        self, analyzer: ContentAnalyzer, simple_keywords: KeywordConfig
    ) -> None:
        """Test that result includes document stats."""
        text = " ".join(["word"] * 100)
        nodes = [create_paragraph_node(text)]
        ast = create_test_ast(text, nodes=nodes)

        result = analyzer.analyze(ast, simple_keywords)

        assert result.document_stats.word_count > 0

    def test_analyze_has_recommendations(
        self, analyzer: ContentAnalyzer, simple_keywords: KeywordConfig
    ) -> None:
        """Test that result includes recommendations."""
        ast = create_test_ast("Short content.")

        result = analyzer.analyze(ast, simple_keywords)

        assert isinstance(result.recommendations, list)


# =============================================================================
# GEO Score Component Tests
# =============================================================================


class TestGEOScoreComponents:
    """Tests for GEO score component calculation."""

    def test_seo_score_calculated(
        self, analyzer: ContentAnalyzer, simple_keywords: KeywordConfig
    ) -> None:
        """Test that SEO score is calculated."""
        nodes = [
            create_heading_node("Cloud Computing Guide", level=1),
            create_paragraph_node("Cloud computing is essential for modern businesses."),
        ]
        ast = create_test_ast(
            "Cloud Computing Guide. Cloud computing is essential.",
            nodes=nodes,
        )

        result = analyzer.analyze(ast, simple_keywords)

        assert result.geo_score.seo_score.total >= 0

    def test_semantic_score_calculated(
        self, analyzer: ContentAnalyzer, simple_keywords: KeywordConfig
    ) -> None:
        """Test that semantic score is calculated."""
        text = "Cloud computing with data center and virtualization."
        nodes = [create_paragraph_node(text)]
        ast = create_test_ast(text, nodes=nodes)

        result = analyzer.analyze(ast, simple_keywords)

        assert result.geo_score.semantic_score.total >= 0

    def test_ai_score_calculated(
        self, analyzer: ContentAnalyzer, simple_keywords: KeywordConfig
    ) -> None:
        """Test that AI score is calculated."""
        text = "Cloud computing enables scalable infrastructure solutions."
        nodes = [create_paragraph_node(text)]
        ast = create_test_ast(text, nodes=nodes)

        result = analyzer.analyze(ast, simple_keywords)

        assert result.geo_score.ai_score.total >= 0

    def test_readability_score_calculated(
        self, analyzer: ContentAnalyzer, simple_keywords: KeywordConfig
    ) -> None:
        """Test that readability score is calculated."""
        text = "Cloud computing is great. It helps businesses grow. Simple and effective."
        nodes = [create_paragraph_node(text)]
        ast = create_test_ast(text, nodes=nodes)

        result = analyzer.analyze(ast, simple_keywords)

        assert result.geo_score.readability_score.total >= 0


# =============================================================================
# File Analysis Tests
# =============================================================================


class TestFileAnalysis:
    """Tests for analyzing DOCX files directly."""

    def test_analyze_file(
        self, analyzer: ContentAnalyzer, temp_docx: Path
    ) -> None:
        """Test analyzing a DOCX file."""
        result = analyzer.analyze_file(
            file_path=temp_docx,
            primary_keyword="cloud computing",
            secondary_keywords=["AWS", "Azure"],
        )

        assert isinstance(result, AnalysisResult)
        assert result.geo_score.total >= 0

    def test_analyze_file_with_all_params(
        self, analyzer: ContentAnalyzer, temp_docx: Path
    ) -> None:
        """Test file analysis with all parameters."""
        result = analyzer.analyze_file(
            file_path=temp_docx,
            primary_keyword="cloud computing",
            secondary_keywords=["AWS", "Azure"],
            semantic_entities=["data center"],
            title="Cloud Computing Guide",
            url_slug="cloud-computing-guide",
            meta_description="A comprehensive guide to cloud computing.",
        )

        assert isinstance(result, AnalysisResult)


# =============================================================================
# Version Comparison Tests
# =============================================================================


class TestVersionComparison:
    """Tests for comparing document versions."""

    def test_compare_versions(
        self, analyzer: ContentAnalyzer, simple_keywords: KeywordConfig
    ) -> None:
        """Test comparing original and optimized versions."""
        # Original (short content)
        original_nodes = [
            create_paragraph_node("Cloud computing is useful."),
        ]
        original_ast = create_test_ast(
            "Cloud computing is useful.",
            nodes=original_nodes,
        )

        # Optimized (expanded content)
        optimized_text = (
            "Cloud computing is useful for modern businesses. "
            "AWS and Azure provide scalable infrastructure. "
            "Data centers enable virtualization and scalability."
        )
        optimized_nodes = [
            create_heading_node("Cloud Computing Guide", level=1),
            create_paragraph_node(optimized_text),
        ]
        optimized_ast = create_test_ast(optimized_text, nodes=optimized_nodes)

        comparison = analyzer.compare_versions(
            original_ast, optimized_ast, simple_keywords
        )

        assert comparison.original_score >= 0
        assert comparison.optimized_score >= 0
        assert isinstance(comparison.improvement, float)

    def test_compare_identifies_improvements(
        self, analyzer: ContentAnalyzer, simple_keywords: KeywordConfig
    ) -> None:
        """Test that comparison identifies improvements."""
        # Original (poor structure)
        original = create_test_ast("Just some text.", nodes=[])

        # Optimized (good structure)
        optimized_nodes = [
            create_heading_node("Cloud Computing", level=1),
            create_paragraph_node("Cloud computing with AWS and Azure."),
            create_heading_node("Benefits", level=2),
            create_paragraph_node("Scalability and data centers."),
        ]
        optimized = create_test_ast(
            "Cloud Computing. Cloud computing with AWS. Benefits. Scalability.",
            nodes=optimized_nodes,
        )

        comparison = analyzer.compare_versions(original, optimized, simple_keywords)

        # Optimized should score better
        assert comparison.optimized_score >= comparison.original_score


# =============================================================================
# Document Stats Tests
# =============================================================================


class TestDocumentStats:
    """Tests for document statistics calculation."""

    def test_word_count_calculated(
        self, analyzer: ContentAnalyzer, simple_keywords: KeywordConfig
    ) -> None:
        """Test word count calculation."""
        text = "one two three four five"
        ast = create_test_ast(text, nodes=[create_paragraph_node(text)])

        result = analyzer.analyze(ast, simple_keywords)

        assert result.document_stats.word_count == 5

    def test_heading_count_calculated(
        self, analyzer: ContentAnalyzer, simple_keywords: KeywordConfig
    ) -> None:
        """Test heading count calculation."""
        nodes = [
            create_heading_node("H1", level=1),
            create_heading_node("H2", level=2),
            create_heading_node("H3", level=3),
        ]
        ast = create_test_ast("H1 H2 H3", nodes=nodes)

        result = analyzer.analyze(ast, simple_keywords)

        assert result.document_stats.heading_count == 3


# =============================================================================
# Configuration Tests
# =============================================================================


class TestAnalyzerConfig:
    """Tests for analyzer configuration."""

    def test_custom_weights(self) -> None:
        """Test custom score weights."""
        config = AnalyzerConfig(
            seo_weight=0.30,
            semantic_weight=0.25,
            ai_weight=0.25,
            readability_weight=0.20,
        )
        analyzer = ContentAnalyzer(config)

        assert analyzer.config.seo_weight == 0.30

    def test_disable_semantic_analysis(self) -> None:
        """Test disabling semantic analysis."""
        config = AnalyzerConfig(enable_semantic_analysis=False)
        analyzer = ContentAnalyzer(config)
        keywords = KeywordConfig(primary_keyword="test")
        ast = create_test_ast("Test content", nodes=[])

        result = analyzer.analyze(ast, keywords)

        # Should still complete analysis
        assert isinstance(result, AnalysisResult)

    def test_disable_ai_analysis(self) -> None:
        """Test disabling AI analysis."""
        config = AnalyzerConfig(enable_ai_analysis=False)
        analyzer = ContentAnalyzer(config)
        keywords = KeywordConfig(primary_keyword="test")
        ast = create_test_ast("Test content", nodes=[])

        result = analyzer.analyze(ast, keywords)

        # AI score should use default
        assert result.geo_score.ai_score.total == 75


# =============================================================================
# Integration Tests
# =============================================================================


class TestFullIntegration:
    """Full integration tests."""

    def test_full_analysis_workflow(
        self, analyzer: ContentAnalyzer, temp_docx: Path
    ) -> None:
        """Test complete analysis workflow."""
        # Analyze file
        result = analyzer.analyze_file(
            file_path=temp_docx,
            primary_keyword="cloud computing",
            secondary_keywords=["AWS", "Azure", "serverless"],
            semantic_entities=["data center", "virtualization"],
        )

        # Verify all components present
        assert result.document_stats.word_count > 0
        assert result.geo_score.total >= 0
        assert result.geo_score.seo_score.total >= 0
        assert result.geo_score.semantic_score.total >= 0
        assert result.geo_score.ai_score.total >= 0
        assert result.geo_score.readability_score.total >= 0
        assert isinstance(result.recommendations, list)

    def test_summary_generation(
        self, analyzer: ContentAnalyzer, simple_keywords: KeywordConfig
    ) -> None:
        """Test analysis summary generation."""
        nodes = [
            create_heading_node("Test", level=1),
            create_paragraph_node("This is test content for cloud computing."),
        ]
        ast = create_test_ast(
            "Test. This is test content for cloud computing.",
            nodes=nodes,
        )

        result = analyzer.analyze(ast, simple_keywords)

        summary = result.summary
        assert "GEO Score" in summary
        assert "SEO" in summary


# =============================================================================
# Convenience Function Tests
# =============================================================================


class TestConvenienceFunctions:
    """Tests for convenience functions."""

    def test_analyze_content_function(self) -> None:
        """Test analyze_content convenience function."""
        nodes = [
            create_heading_node("Python Guide", level=1),
            create_paragraph_node("Learn Python programming."),
        ]
        ast = create_test_ast(
            "Python Guide. Learn Python programming.",
            nodes=nodes,
        )

        result = analyze_content(
            ast=ast,
            primary_keyword="python",
            secondary_keywords=["programming"],
        )

        assert isinstance(result, AnalysisResult)
        assert result.geo_score.seo_score.keyword_analysis.primary_found is True

    def test_analyze_docx_function(self, temp_docx: Path) -> None:
        """Test analyze_docx convenience function."""
        result = analyze_docx(
            file_path=temp_docx,
            primary_keyword="cloud computing",
        )

        assert isinstance(result, AnalysisResult)
        assert result.document_stats.word_count > 0
